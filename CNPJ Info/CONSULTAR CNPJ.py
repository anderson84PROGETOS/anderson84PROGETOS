import tkinter as tk
from tkinter import scrolledtext, messagebox, filedialog, OptionMenu
import requests
import webbrowser
from datetime import datetime
import openpyxl
from openpyxl.utils import get_column_letter
from docx import Document

def consultar_cnpj(cnpj):
    url = f'https://www.receitaws.com.br/v1/cnpj/{cnpj}'
    try:
        response = requests.get(url)
        if response.status_code == 200:
            return response.json()
        else:
            messagebox.showerror("Erro ao consultar CNPJ", f"Erro ao consultar CNPJ: {response.status_code}")
    except requests.exceptions.RequestException as e:
        messagebox.showerror("Erro na requisição", f"Erro na requisição: {e}")

def calcular_idade(data_abertura):
    hoje = datetime.now()
    data_abertura = datetime.strptime(data_abertura, '%d/%m/%Y')
    diferenca = hoje - data_abertura
    anos = diferenca.days // 365
    meses = (diferenca.days % 365) // 30
    dias = (diferenca.days % 365) % 30
    return f"{anos} anos, {meses} meses e {dias} dias"

def abrir_google_maps(logradouro, numero, municipio, uf):
    endereco = f"{logradouro}, {numero}, {municipio}, {uf}"
    url = f"https://www.google.com/maps/search/?api=1&query={endereco.replace(' ', '+')}"
    webbrowser.open(url)

def limpar_cnpj(cnpj):
    return ''.join(filter(str.isdigit, cnpj))

def gerar_texto(dados_cnpj):
    logradouro = dados_cnpj.get('logradouro', 'Não encontrado')
    numero = dados_cnpj.get('numero', 'Não encontrado')
    municipio = dados_cnpj.get('municipio', 'Não encontrado')
    uf = dados_cnpj.get('uf', 'Não encontrado')

    message = f"""
CNPJ: {dados_cnpj.get('cnpj', 'Não encontrado')}

RAZÃO SOCIAL: {dados_cnpj.get('nome', 'Não encontrado')}

MATRIZ OU FILIAL: {dados_cnpj.get('tipo', 'Não encontrado')}

NOME FANTASIA: {dados_cnpj.get('fantasia', 'Não encontrado')}

SITUAÇÃO CADASTRAL: {dados_cnpj.get('situacao', 'Não encontrado')}

DATA DA SITUAÇÃO CADASTRAL: {dados_cnpj.get('data_situacao', 'Não encontrado')}

MOTIVO DA SITUAÇÃO CADASTRAL: {dados_cnpj.get('motivo_situacao', 'Não encontrado')}

NATUREZA JURÍDICA: {dados_cnpj.get('natureza_juridica', 'Não encontrado')}

DATA DE ABERTURA: {dados_cnpj.get('abertura', 'Não encontrado')}

IDADE: {calcular_idade(dados_cnpj['abertura']) if 'abertura' in dados_cnpj else 'Não encontrado'}

PORTE (RFB): {dados_cnpj.get('porte', 'Não encontrado')}

CAPITAL SOCIAL: R$ {dados_cnpj.get('capital_social', 'Não encontrado')}

ATUALIZAÇÃO DESTA PÁGINA: {dados_cnpj.get('ultima_atualizacao', 'Não encontrado')}

LOCALIZAÇÃO
=============
ENDEREÇO: {logradouro}     | Número: {numero}
COMPLEMENTO: {dados_cnpj.get('complemento', 'Não encontrado')}
BAIRRO: {dados_cnpj.get('bairro', 'Não encontrado')}
CIDADE | ESTADO: {municipio} | {uf}
CEP: {dados_cnpj.get('cep', 'Não encontrado')}
TELEFONES: {dados_cnpj.get('telefone', 'Não encontrado')}
E-MAILS: {dados_cnpj.get('email', 'Não encontrado')}

ATIVIDADE ECONÔMICA PRINCIPAL
==============================
CÓDIGO: {dados_cnpj['atividade_principal'][0]['code'] if 'atividade_principal' in dados_cnpj else 'Não encontrado'}
DESCRIÇÃO: {dados_cnpj['atividade_principal'][0]['text'] if 'atividade_principal' in dados_cnpj else 'Não encontrado'}

ATIVIDADES ECONÔMICAS SECUNDÁRIAS
====================================
"""
    if 'atividades_secundarias' in dados_cnpj:
        for atividade in dados_cnpj['atividades_secundarias']:
            message += f"CÓDIGO: {atividade['code']} | DESCRIÇÃO: {atividade['text']}\n"

    message += "\nQUADRO DE SÓCIOS E ADMINISTRADORES (QSA)\n==========================================\n"
    if 'qsa' in dados_cnpj:
        for socio in dados_cnpj['qsa']:
            data_entrada = socio.get('data_entrada', None)
            if data_entrada:
                data_entrada = datetime.strptime(data_entrada, '%d/%m/%Y').strftime('%d/%m/%Y')
                message += f"""
NOME: {socio.get('nome', 'Não encontrado')}
QUALIFICAÇÃO: {socio.get('qual', 'Não encontrado')}
ENTRADA: {data_entrada}
"""
            else:
                message += f"""
NOME: {socio.get('nome', 'Não encontrado')}
QUALIFICAÇÃO: {socio.get('qual', 'Não encontrado')}
"""
    else:
        message += "Não encontrado\n"
    
    return message

def salvar_xlsx(dados_cnpj, filename):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Consulta CNPJ"

    headers = ["Categoria", "Informação"]
    ws.append(headers)

    ws.append(["Informações Gerais", ""])
    ws.append(["CNPJ", dados_cnpj.get('cnpj', 'Não encontrado')])
    ws.append(["RAZÃO SOCIAL", dados_cnpj.get('nome', 'Não encontrado')])
    ws.append(["MATRIZ OU FILIAL", dados_cnpj.get('tipo', 'Não encontrado')])
    ws.append(["NOME FANTASIA", dados_cnpj.get('fantasia', 'Não encontrado')])
    ws.append(["SITUAÇÃO CADASTRAL", dados_cnpj.get('situacao', 'Não encontrado')])
    ws.append(["DATA DA SITUAÇÃO CADASTRAL", dados_cnpj.get('data_situacao', 'Não encontrado')])
    ws.append(["MOTIVO DA SITUAÇÃO CADASTRAL", dados_cnpj.get('motivo_situacao', 'Não encontrado')])
    ws.append(["NATUREZA JURÍDICA", dados_cnpj.get('natureza_juridica', 'Não encontrado')])
    ws.append(["DATA DE ABERTURA", dados_cnpj.get('abertura', 'Não encontrado')])
    ws.append(["IDADE", calcular_idade(dados_cnpj['abertura']) if 'abertura' in dados_cnpj else 'Não encontrado'])    
    ws.append(["PORTE (RFB)", dados_cnpj.get('porte', 'Não encontrado')])
    ws.append(["CAPITAL SOCIAL", f"R$ {dados_cnpj.get('capital_social', 'Não encontrado')}"])
    ws.append(["ATUALIZAÇÃO DESTA PÁGINA", dados_cnpj.get('ultima_atualizacao', 'Não encontrado')])
    ws.append(["", ""])
    ws.append(["Localização", ""])
    ws.append(["ENDEREÇO", f"{dados_cnpj.get('logradouro', 'Não encontrado')} | Número: {dados_cnpj.get('numero', 'Não encontrado')}"])
    ws.append(["COMPLEMENTO", dados_cnpj.get('complemento', 'Não encontrado')])
    ws.append(["BAIRRO", dados_cnpj.get('bairro', 'Não encontrado')])
    ws.append(["CIDADE | ESTADO", f"{dados_cnpj.get('municipio', 'Não encontrado')} | {dados_cnpj.get('uf', 'Não encontrado')}"])
    ws.append(["CEP", dados_cnpj.get('cep', 'Não encontrado')])
    ws.append(["TELEFONES", dados_cnpj.get('telefone', 'Não encontrado')])
    ws.append(["E-MAILS", dados_cnpj.get('email', 'Não encontrado')])
    ws.append(["", ""])
    ws.append(["Atividade Econômica Principal", ""])
    ws.append(["CÓDIGO", dados_cnpj['atividade_principal'][0]['code'] if 'atividade_principal' in dados_cnpj else 'Não encontrado'])
    ws.append(["DESCRIÇÃO", dados_cnpj['atividade_principal'][0]['text'] if 'atividade_principal' in dados_cnpj else 'Não encontrado'])
    ws.append(["", ""])
    ws.append(["Atividades Econômicas Secundárias", ""])
    if 'atividades_secundarias' in dados_cnpj:
        for atividade in dados_cnpj['atividades_secundarias']:
            ws.append(["CÓDIGO | DESCRIÇÃO", f"{atividade['code']} | {atividade['text']}"])
    else:
        ws.append(["", "Não encontrado"])

    ws.append(["", ""])
    ws.append(["Quadro de Sócios e Administradores (QSA)", ""])
    if 'qsa' in dados_cnpj:
        for socio in dados_cnpj['qsa']:
            data_entrada = socio.get('data_entrada', None)
            ws.append(["NOME", socio.get('nome', 'Não encontrado')])
            ws.append(["QUALIFICAÇÃO", socio.get('qual', 'Não encontrado')])
            if data_entrada:
                data_entrada = datetime.strptime(data_entrada, '%d/%m/%Y').strftime('%d/%m/%Y')
                ws.append(["ENTRADA", data_entrada])
    else:
        ws.append(["", "Não encontrado"])

    for col in range(1, ws.max_column + 1):
        ws.column_dimensions[get_column_letter(col)].width = 50

    wb.save(filename)

def salvar_docx(dados_cnpj, filename):
    doc = Document()
    doc.add_heading('Consulta CNPJ', 0)

    doc.add_heading('Informações Gerais', level=1)
    doc.add_paragraph(f"CNPJ: {dados_cnpj.get('cnpj', 'Não encontrado')}")
    doc.add_paragraph(f"RAZÃO SOCIAL: {dados_cnpj.get('nome', 'Não encontrado')}")
    doc.add_paragraph(f"MATRIZ OU FILIAL: {dados_cnpj.get('tipo', 'Não encontrado')}")
    doc.add_paragraph(f"NOME FANTASIA: {dados_cnpj.get('fantasia', 'Não encontrado')}")
    doc.add_paragraph(f"SITUAÇÃO CADASTRAL: {dados_cnpj.get('situacao', 'Não encontrado')}")
    doc.add_paragraph(f"DATA DA SITUAÇÃO CADASTRAL: {dados_cnpj.get('data_situacao', 'Não encontrado')}")
    doc.add_paragraph(f"MOTIVO DA SITUAÇÃO CADASTRAL: {dados_cnpj.get('motivo_situacao', 'Não encontrado')}")
    doc.add_paragraph(f"NATUREZA JURÍDICA: {dados_cnpj.get('natureza_juridica', 'Não encontrado')}")
    doc.add_paragraph(f"DATA DE ABERTURA: {dados_cnpj.get('abertura', 'Não encontrado')}")
    doc.add_paragraph(f"IDADE: {calcular_idade(dados_cnpj['abertura']) if 'abertura' in dados_cnpj else 'Não encontrado'}")
    doc.add_paragraph(f"PORTE (RFB): {dados_cnpj.get('porte', 'Não encontrado')}")
    doc.add_paragraph(f"CAPITAL SOCIAL: R$ {dados_cnpj.get('capital_social', 'Não encontrado')}")
    doc.add_paragraph(f"ATUALIZAÇÃO DESTA PÁGINA: {dados_cnpj.get('ultima_atualizacao', 'Não encontrado')}")
    doc.add_heading('Localização', level=1)
    doc.add_paragraph(f"ENDEREÇO: {dados_cnpj.get('logradouro', 'Não encontrado')} | Número: {dados_cnpj.get('numero', 'Não encontrado')}")
    doc.add_paragraph(f"COMPLEMENTO: {dados_cnpj.get('complemento', 'Não encontrado')}")
    doc.add_paragraph(f"BAIRRO: {dados_cnpj.get('bairro', 'Não encontrado')}")
    doc.add_paragraph(f"CIDADE | ESTADO: {dados_cnpj.get('municipio', 'Não encontrado')} | {dados_cnpj.get('uf', 'Não encontrado')}")
    doc.add_paragraph(f"CEP: {dados_cnpj.get('cep', 'Não encontrado')}")
    doc.add_paragraph(f"TELEFONES: {dados_cnpj.get('telefone', 'Não encontrado')}")
    doc.add_paragraph(f"E-MAILS: {dados_cnpj.get('email', 'Não encontrado')}")
    doc.add_heading('Atividade Econômica Principal', level=1)
    doc.add_paragraph(f"CÓDIGO: {dados_cnpj['atividade_principal'][0]['code'] if 'atividade_principal' in dados_cnpj else 'Não encontrado'}")
    doc.add_paragraph(f"DESCRIÇÃO: {dados_cnpj['atividade_principal'][0]['text'] if 'atividade_principal' in dados_cnpj else 'Não encontrado'}")
    doc.add_heading('Atividades Econômicas Secundárias', level=1)
    if 'atividades_secundarias' in dados_cnpj:
        for atividade in dados_cnpj['atividades_secundarias']:
            doc.add_paragraph(f"CÓDIGO: {atividade['code']} | DESCRIÇÃO: {atividade['text']}")
    else:
        doc.add_paragraph("Não encontrado")

    doc.add_heading('Quadro de Sócios e Administradores (QSA)', level=1)
    if 'qsa' in dados_cnpj:
        for socio in dados_cnpj['qsa']:
            data_entrada = socio.get('data_entrada', None)
            if data_entrada:
                data_entrada = datetime.strptime(data_entrada, '%d/%m/%Y').strftime('%d/%m/%Y')
                doc.add_paragraph(f"NOME: {socio.get('nome', 'Não encontrado')}")
                doc.add_paragraph(f"QUALIFICAÇÃO: {socio.get('qual', 'Não encontrado')}")
                doc.add_paragraph(f"ENTRADA: {data_entrada}")
            else:
                doc.add_paragraph(f"NOME: {socio.get('nome', 'Não encontrado')}")
                doc.add_paragraph(f"QUALIFICAÇÃO: {socio.get('qual', 'Não encontrado')}")
    else:
        doc.add_paragraph("Não encontrado")

    doc.save(filename)

def salvar_txt(dados_cnpj, filename):
    with open(filename, 'w', encoding='utf-8') as f:
        f.write(gerar_texto(dados_cnpj))

def salvar_arquivo(dados_cnpj, formato_var):
    formato = formato_var.get()
    cnpj = dados_cnpj.get('cnpj', 'desconhecido').replace('/', '_').replace('.', '_').replace('-', '_')
    
    if formato == "Excel (.xlsx)":
        default_extension = ".xlsx"
        filetypes = [("Planilhas Excel", "*.xlsx")]
        default_filename = f"consulta_cnpj_{cnpj}.xlsx"
    elif formato == "Word (.docx)":
        default_extension = ".docx"
        filetypes = [("Documentos Word", "*.docx")]
        default_filename = f"consulta_cnpj_{cnpj}.docx"
    else:  # Texto (.txt)
        default_extension = ".txt"
        filetypes = [("Arquivos de Texto", "*.txt")]
        default_filename = f"consulta_cnpj_{cnpj}.txt"

    filename = filedialog.asksaveasfilename(
        defaultextension=default_extension,
        filetypes=filetypes,
        initialfile=default_filename
    )

    if filename:
        try:
            if formato == "Excel (.xlsx)":
                salvar_xlsx(dados_cnpj, filename)
            elif formato == "Word (.docx)":
                salvar_docx(dados_cnpj, filename)
            else:  # Texto (.txt)
                salvar_txt(dados_cnpj, filename)
            messagebox.showinfo("Sucesso", f"Arquivo salvo como {filename}")
        except Exception as e:
            messagebox.showerror("Erro ao salvar", f"Não foi possível salvar o arquivo: {e}")

def consultar_e_mostrar(cnpj_entry, info_text, maps_button, salvar_button, formato_var):
    cnpj = limpar_cnpj(cnpj_entry.get())
    dados_cnpj = consultar_cnpj(cnpj)
    if dados_cnpj:
        info_text.config(state=tk.NORMAL)
        info_text.delete(1.0, tk.END)
        
        message = gerar_texto(dados_cnpj)
        info_text.insert(tk.END, message)
        info_text.config(state=tk.DISABLED)

        logradouro = dados_cnpj.get('logradouro', 'Não encontrado')
        numero = dados_cnpj.get('numero', 'Não encontrado')
        municipio = dados_cnpj.get('municipio', 'Não encontrado')
        uf = dados_cnpj.get('uf', 'Não encontrado')

        maps_button.config(state=tk.NORMAL, command=lambda: abrir_google_maps(logradouro, numero, municipio, uf))
        salvar_button.config(state=tk.NORMAL, command=lambda: salvar_arquivo(dados_cnpj, formato_var))

def abrir_site():
    webbrowser.open("https://www.informecadastral.com.br")

def criar_interface_grafica():
    root = tk.Tk()
    root.wm_state('zoomed')
    root.geometry("1200x950")
    root.title("CONSULTAR CNPJ")

    label = tk.Label(root, text="Digite o número do CNPJ (exemplo: 1823612001  ou  60.701.190/0001-04)", font=("Arial", 11))
    label.pack(pady=5)

    cnpj_entry = tk.Entry(root, width=40, font=("Arial", 11))
    cnpj_entry.pack()

    consultar_button = tk.Button(root, text="Consultar", font=("Arial", 11), bg="#0bfc03")
    consultar_button.pack(pady=5)

    maps_button = tk.Button(root, text="Abrir no Google Maps", font=("Arial", 11), bg="#00FFFF", state=tk.DISABLED)
    maps_button.pack(pady=5)

    formato_var = tk.StringVar(root)
    formato_var.set("Excel (.xlsx)")  # Valor padrão
    formato_menu = OptionMenu(root, formato_var, "Excel (.xlsx)", "Word (.docx)", "Texto (.txt)")
    formato_menu.config(font=("Arial", 11))
    formato_menu.pack(pady=5)

    salvar_button = tk.Button(root, text="Salvar Arquivo", font=("Arial", 11), bg="#FF4500", state=tk.DISABLED)
    salvar_button.pack(pady=5)

    site_button = tk.Button(root, text="Mais informações no site: https://www.informecadastral.com.br", command=abrir_site, font=("Arial", 11), bg="#ffbf00")
    site_button.pack(pady=5)

    info_text = scrolledtext.ScrolledText(root, wrap=tk.WORD, width=120, height=35, font=("Arial", 12))
    info_text.pack(pady=5)
    info_text.config(state=tk.DISABLED)

    consultar_button.config(command=lambda: consultar_e_mostrar(cnpj_entry, info_text, maps_button, salvar_button, formato_var))

    root.mainloop()

criar_interface_grafica()
