import tkinter as tk
from tkinter import scrolledtext, messagebox, filedialog, ttk
import requests
import webbrowser
from datetime import datetime
import openpyxl
from openpyxl.utils import get_column_letter
from docx import Document
import csv
import json
import xml.etree.ElementTree as ET
from xml.dom import minidom
import os

# ==================== ESTILO ====================
class HackerStyle:
    BG = "#0a0a0a"
    FG = "#00ff41"
    ACCENT = "#00ff41"
    RED = "#ff0033"
    GRAY = "#1a1a1a"
    FONT = ("Consolas", 11)
    FONT_BOLD = ("Consolas", 12, "bold")
    TITLE_FONT = ("Consolas", 18, "bold")

def abrir_arquivo():
    caminho = filedialog.askopenfilename(
        title="Abrir arquivo",
        filetypes=[
            ("Todos os suportados", "*.txt *.docx *.xlsx *.csv *.json *.log *.md *.html *.htm *.xml"),
            ("Texto", "*.txt *.log *.md"),
            ("Word", "*.docx"),
            ("Excel", "*.xlsx"),
            ("CSV", "*.csv"),
            ("JSON", "*.json"),
            ("HTML", "*.html *.htm"),
            ("XML", "*.xml"),
            ("Todos os arquivos", "*.*")
        ]
    )
    if not caminho:
        return

    try:
        ext = os.path.splitext(caminho)[1].lower()
        conteudo = ""

        if ext in [".txt", ".log", ".md", ".html", ".htm"]:
            with open(caminho, "r", encoding="utf-8") as f:
                conteudo = f.read()

        elif ext == ".xml":
            try:
                tree = ET.parse(caminho)
                root = tree.getroot()
                conteudo = minidom.parse(caminho).toprettyxml(indent="  ")
            except ET.ParseError:
                with open(caminho, "r", encoding="utf-8") as f:
                    conteudo = f.read()

        elif ext == ".json":
            with open(caminho, "r", encoding="utf-8") as f:
                try:
                    dados = json.load(f)
                    conteudo = json.dumps(dados, indent=4, ensure_ascii=False)
                except json.JSONDecodeError:
                    conteudo = "Erro: Arquivo JSON malformado"

        elif ext == ".csv":
            with open(caminho, "r", encoding="utf-8") as f:
                try:
                    linhas = csv.reader(f)
                    conteudo = "\n".join([",".join(row) for row in linhas])
                except csv.Error:
                    conteudo = "Erro: Arquivo CSV malformado"

        elif ext == ".docx":
            try:
                doc = Document(caminho)
                conteudo = "\n".join([p.text for p in doc.paragraphs if p.text])
            except Exception as e:
                conteudo = f"Erro ao ler arquivo DOCX: {str(e)}"

        elif ext == ".xlsx":
            try:
                wb = openpyxl.load_workbook(caminho)
                sheet = wb.active
                linhas = []
                for row in sheet.iter_rows(values_only=True):
                    linha = "\t".join([str(cell) if cell is not None else "" for cell in row])
                    linhas.append(linha)
                conteudo = "\n".join(linhas)
            except Exception as e:
                conteudo = f"Erro ao ler arquivo XLSX: {str(e)}"

        else:
            messagebox.showerror("Erro", f"Tipo de arquivo não suportado: {ext}")
            return

        text_area.config(state=tk.NORMAL)
        text_area.delete(1.0, tk.END)
        text_area.insert(tk.END, conteudo)
        text_area.config(state=tk.DISABLED)

    except Exception as e:
        messagebox.showerror("Erro ao abrir o arquivo", f"Erro: {str(e)}")

def consultar_cnpj(cnpj):
    url = f'https://www.receitaws.com.br/v1/cnpj/{cnpj}'
    try:
        response = requests.get(url)
        if response.status_code == 200:
            return response.json()
        else:
            messagebox.showerror("Erro ao consultar CNPJ", f"Erro ao consultar CNPJ: {response.status_code}")
    except requests.exceptions.RequestException as e:
        messagebox.showerror("Erro na requisição", f"Erro na requisição: {e}")
    return None

def calcular_idade(data_abertura):
    try:
        hoje = datetime.now()
        data_abertura = datetime.strptime(data_abertura, '%d/%m/%Y')
        diferenca = hoje - data_abertura
        anos = diferenca.days // 365
        meses = (diferenca.days % 365) // 30
        dias = (diferenca.days % 365) % 30
        return f"{anos} anos, {meses} meses e {dias} dias"
    except ValueError:
        return "Data de abertura inválida"

def abrir_google_maps(logradouro, numero, municipio, uf):
    endereco = f"{logradouro}, {numero}, {municipio}, {uf}"
    url = f"https://www.google.com/maps/search/?api=1&query={endereco.replace(' ', '+')}"
    webbrowser.open(url)    

def limpar_cnpj(cnpj):
    return ''.join(filter(str.isdigit, cnpj))

def gerar_texto(dados_cnpj):
    logradouro = dados_cnpj.get('logradouro', 'Não encontrado')
    numero = dados_cnpj.get('numero', 'Não encontrado')
    municipio = dados_cnpj.get('municipio', 'Não encontrado')
    uf = dados_cnpj.get('uf', 'Não encontrado')

    message = f"""
CNPJ: {dados_cnpj.get('cnpj', 'Não encontrado')}

RAZÃO SOCIAL: {dados_cnpj.get('nome', 'Não encontrado')}

MATRIZ OU FILIAL: {dados_cnpj.get('tipo', 'Não encontrado')}

NOME FANTASIA: {dados_cnpj.get('fantasia', 'Não encontrado')}

SITUAÇÃO CADASTRAL: {dados_cnpj.get('situacao', 'Não encontrado')}

DATA DA SITUAÇÃO CADASTRAL: {dados_cnpj.get('data_situacao', 'Não encontrado')}

MOTIVO DA SITUAÇÃO CADASTRAL: {dados_cnpj.get('motivo_situacao', 'Não encontrado')}

NATUREZA JURÍDICA: {dados_cnpj.get('natureza_juridica', 'Não encontrado')}

DATA DE ABERTURA: {dados_cnpj.get('abertura', 'Não encontrado')}

IDADE: {calcular_idade(dados_cnpj['abertura']) if 'abertura' in dados_cnpj else 'Não encontrado'}

PORTE (RFB): {dados_cnpj.get('porte', 'Não encontrado')}

CAPITAL SOCIAL: R$ {dados_cnpj.get('capital_social', 'Não encontrado')}

ATUALIZAÇÃO DESTA PÁGINA: {dados_cnpj.get('ultima_atualizacao', 'Não encontrado')}



LOCALIZAÇÃO
===========

ENDEREÇO: {logradouro}       |  Número: {numero}

COMPLEMENTO: {dados_cnpj.get('complemento', 'Não encontrado')}

BAIRRO: {dados_cnpj.get('bairro', 'Não encontrado')}

CIDADE | ESTADO: {municipio} | {uf}

CEP: {dados_cnpj.get('cep', 'Não encontrado')}

TELEFONES: {dados_cnpj.get('telefone', 'Não encontrado')}

E-MAILS: {dados_cnpj.get('email', 'Não encontrado')}



ATIVIDADE ECONÔMICA PRINCIPAL
==============================

CÓDIGO: {dados_cnpj['atividade_principal'][0]['code'] if 'atividade_principal' in dados_cnpj and dados_cnpj['atividade_principal'] else 'Não encontrado'}
DESCRIÇÃO: {dados_cnpj['atividade_principal'][0]['text'] if 'atividade_principal' in dados_cnpj and dados_cnpj['atividade_principal'] else 'Não encontrado'}


ATIVIDADES ECONÔMICAS SECUNDÁRIAS
=================================

"""
    if 'atividades_secundarias' in dados_cnpj and dados_cnpj['atividades_secundarias']:
        for atividade in dados_cnpj['atividades_secundarias']:
            message += f"CÓDIGO: {atividade['code']} | DESCRIÇÃO: {atividade['text']}\n"
    else:
        message += "Não encontrado\n"

    message += "\n\nQUADRO DE SÓCIOS E ADMINISTRADORES (QSA)\n==========================================\n"
    if 'qsa' in dados_cnpj and dados_cnpj['qsa']:
        for socio in dados_cnpj['qsa']:
            data_entrada = socio.get('data_entrada', None)
            if data_entrada:
                try:
                    data_entrada = datetime.strptime(data_entrada, '%d/%m/%Y').strftime('%d/%m/%Y')
                    message += f"""
NOME: {socio.get('nome', 'Não encontrado')}
QUALIFICAÇÃO: {socio.get('qual', 'Não encontrado')}
ENTRADA: {data_entrada}
"""
                except ValueError:
                    message += f"""
NOME: {socio.get('nome', 'Não encontrado')}
QUALIFICAÇÃO: {socio.get('qual', 'Não encontrado')}
ENTRADA: Data inválida
"""
            else:
                message += f"""
NOME: {socio.get('nome', 'Não encontrado')}
QUALIFICAÇÃO: {socio.get('qual', 'Não encontrado')}
"""
    else:
        message += "Não encontrado\n"
    
    return message

def salvar_xlsx(dados_cnpj, filename):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Consulta CNPJ"

    headers = ["Categoria", "Informação\n"]

    ws.append(headers)

    ws.append(["Informações Gerais\n", ""])

    ws.append(["CNPJ", dados_cnpj.get('cnpj', 'Não encontrado')])
    ws.append(["RAZÃO SOCIAL", dados_cnpj.get('nome', 'Não encontrado')])
    ws.append(["MATRIZ OU FILIAL", dados_cnpj.get('tipo', 'Não encontrado')])
    ws.append(["NOME FANTASIA", dados_cnpj.get('fantasia', 'Não encontrado')])
    ws.append(["SITUAÇÃO CADASTRAL", dados_cnpj.get('situacao', 'Não encontrado')])
    ws.append(["DATA DA SITUAÇÃO CADASTRAL", dados_cnpj.get('data_situacao', 'Não encontrado')])
    ws.append(["MOTIVO DA SITUAÇÃO CADASTRAL", dados_cnpj.get('motivo_situacao', 'Não encontrado')])
    ws.append(["NATUREZA JURÍDICA", dados_cnpj.get('natureza_juridica', 'Não encontrado')])
    ws.append(["DATA DE ABERTURA", dados_cnpj.get('abertura', 'Não encontrado')])
    ws.append(["IDADE", calcular_idade(dados_cnpj['abertura']) if 'abertura' in dados_cnpj else 'Não encontrado'])
    ws.append(["PORTE (RFB)", dados_cnpj.get('porte', 'Não encontrado')])
    ws.append(["CAPITAL SOCIAL", f"R$ {dados_cnpj.get('capital_social', 'Não encontrado')}"])
    ws.append(["ATUALIZAÇÃO DESTA PÁGINA", dados_cnpj.get('ultima_atualizacao', 'Não encontrado')])
    ws.append(["", ""])

    ws.append(["\nLocalização\n", ""])

    ws.append(["ENDEREÇO", f"{dados_cnpj.get('logradouro', 'Não encontrado')} | Número: {dados_cnpj.get('numero', 'Não encontrado')}"])
    ws.append(["COMPLEMENTO", dados_cnpj.get('complemento', 'Não encontrado')])
    ws.append(["BAIRRO", dados_cnpj.get('bairro', 'Não encontrado')])
    ws.append(["CIDADE | ESTADO", f"{dados_cnpj.get('municipio', 'Não encontrado')} | {dados_cnpj.get('uf', 'Não encontrado')}"])
    ws.append(["CEP", dados_cnpj.get('cep', 'Não encontrado')])
    ws.append(["TELEFONES", dados_cnpj.get('telefone', 'Não encontrado')])
    ws.append(["E-MAILS", dados_cnpj.get('email', 'Não encontrado')])
    ws.append(["", ""])

    ws.append(["Atividade Econômica Principal", ""])
    ws.append(["CÓDIGO", dados_cnpj['atividade_principal'][0]['code'] if 'atividade_principal' in dados_cnpj and dados_cnpj['atividade_principal'] else 'Não encontrado'])
    ws.append(["DESCRIÇÃO", dados_cnpj['atividade_principal'][0]['text'] if 'atividade_principal' in dados_cnpj and dados_cnpj['atividade_principal'] else 'Não encontrado'])
    ws.append(["", ""])

    ws.append(["Atividades Econômicas Secundárias", ""])
    if 'atividades_secundarias' in dados_cnpj and dados_cnpj['atividades_secundarias']:
        for atividade in dados_cnpj['atividades_secundarias']:
            ws.append(["CÓDIGO | DESCRIÇÃO", f"{atividade['code']} | {atividade['text']}"])
    else:
        ws.append(["", "Não encontrado"])

    ws.append(["", ""])
    ws.append(["Quadro de Sócios e Administradores (QSA)", ""])
    if 'qsa' in dados_cnpj and dados_cnpj['qsa']:
        for socio in dados_cnpj['qsa']:
            ws.append(["NOME", socio.get('nome', 'Não encontrado')])
            ws.append(["QUALIFICAÇÃO", socio.get('qual', 'Não encontrado')])
            data_entrada = socio.get('data_entrada', None)
            if data_entrada:
                try:
                    data_entrada = datetime.strptime(data_entrada, '%d/%m/%Y').strftime('%d/%m/%Y')
                    ws.append(["ENTRADA", data_entrada])
                except ValueError:
                    ws.append(["ENTRADA", "Data inválida"])
    else:
        ws.append(["", "Não encontrado"])

    for col in range(1, ws.max_column + 1):
        ws.column_dimensions[get_column_letter(col)].width = 50

    wb.save(filename)

def salvar_docx(dados_cnpj, filename):
    doc = Document()
    doc.add_heading('Consulta CNPJ', 0)

    doc.add_heading('Informações Gerais', level=1)

    doc.add_paragraph(f"CNPJ: {dados_cnpj.get('cnpj', 'Não encontrado')}")
    doc.add_paragraph(f"RAZÃO SOCIAL: {dados_cnpj.get('nome', 'Não encontrado')}")
    doc.add_paragraph(f"MATRIZ OU FILIAL: {dados_cnpj.get('tipo', 'Não encontrado')}")
    doc.add_paragraph(f"NOME FANTASIA: {dados_cnpj.get('fantasia', 'Não encontrado')}")
    doc.add_paragraph(f"SITUAÇÃO CADASTRAL: {dados_cnpj.get('situacao', 'Não encontrado')}")
    doc.add_paragraph(f"DATA DA SITUAÇÃO CADASTRAL: {dados_cnpj.get('data_situacao', 'Não encontrado')}")
    doc.add_paragraph(f"MOTIVO DA SITUAÇÃO CADASTRAL: {dados_cnpj.get('motivo_situacao', 'Não encontrado')}")
    doc.add_paragraph(f"NATUREZA JURÍDICA: {dados_cnpj.get('natureza_juridica', 'Não encontrado')}")
    doc.add_paragraph(f"DATA DE ABERTURA: {dados_cnpj.get('abertura', 'Não encontrado')}")
    doc.add_paragraph(f"IDADE: {calcular_idade(dados_cnpj['abertura']) if 'abertura' in dados_cnpj else 'Não encontrado'}")
    doc.add_paragraph(f"PORTE (RFB): {dados_cnpj.get('porte', 'Não encontrado')}")
    doc.add_paragraph(f"CAPITAL SOCIAL: R$ {dados_cnpj.get('capital_social', 'Não encontrado')}")
    doc.add_paragraph(f"ATUALIZAÇÃO DESTA PÁGINA: {dados_cnpj.get('ultima_atualizacao', 'Não encontrado')}")

    doc.add_heading('\nLocalização\n', level=1)

    doc.add_paragraph(f"ENDEREÇO: {dados_cnpj.get('logradouro', 'Não encontrado')} | Número: {dados_cnpj.get('numero', 'Não encontrado')}")
    doc.add_paragraph(f"COMPLEMENTO: {dados_cnpj.get('complemento', 'Não encontrado')}")
    doc.add_paragraph(f"BAIRRO: {dados_cnpj.get('bairro', 'Não encontrado')}")
    doc.add_paragraph(f"CIDADE | ESTADO: {dados_cnpj.get('municipio', 'Não encontrado')} | {dados_cnpj.get('uf', 'Não encontrado')}")
    doc.add_paragraph(f"CEP: {dados_cnpj.get('cep', 'Não encontrado')}")
    doc.add_paragraph(f"TELEFONES: {dados_cnpj.get('telefone', 'Não encontrado')}")
    doc.add_paragraph(f"E-MAILS: {dados_cnpj.get('email', 'Não encontrado')}")
    doc.add_heading('Atividade Econômica Principal', level=1)
    doc.add_paragraph(f"CÓDIGO: {dados_cnpj['atividade_principal'][0]['code'] if 'atividade_principal' in dados_cnpj and dados_cnpj['atividade_principal'] else 'Não encontrado'}")
    doc.add_paragraph(f"DESCRIÇÃO: {dados_cnpj['atividade_principal'][0]['text'] if 'atividade_principal' in dados_cnpj and dados_cnpj['atividade_principal'] else 'Não encontrado'}")

    doc.add_heading('Atividades Econômicas Secundárias', level=1)
    if 'atividades_secundarias' in dados_cnpj and dados_cnpj['atividades_secundarias']:
        for atividade in dados_cnpj['atividades_secundarias']:
            doc.add_paragraph(f"CÓDIGO: {atividade['code']} | DESCRIÇÃO: {atividade['text']}")
    else:
        doc.add_paragraph("Não encontrado")

    doc.add_heading('Quadro de Sócios e Administradores (QSA)', level=1)
    if 'qsa' in dados_cnpj and dados_cnpj['qsa']:
        for socio in dados_cnpj['qsa']:
            doc.add_paragraph(f"NOME: {socio.get('nome', 'Não encontrado')}")
            doc.add_paragraph(f"QUALIFICAÇÃO: {socio.get('qual', 'Não encontrado')}")
            data_entrada = socio.get('data_entrada', None)
            if data_entrada:
                try:
                    data_entrada = datetime.strptime(data_entrada, '%d/%m/%Y').strftime('%d/%m/%Y')
                    doc.add_paragraph(f"ENTRADA: {data_entrada}")
                except ValueError:
                    doc.add_paragraph("ENTRADA: Data inválida")
    else:
        doc.add_paragraph("Não encontrado")

    doc.save(filename)

def salvar_txt(dados_cnpj, filename):
    with open(filename, 'w', encoding='utf-8') as f:
        f.write(gerar_texto(dados_cnpj))

def salvar_csv(dados_cnpj, filename):
    with open(filename, 'w', encoding='utf-8', newline='') as f:
        writer = csv.writer(f)
        writer.writerow(["Categoria", "Informação\n"])

        writer.writerow(["Informações Gerais\n", ""])

        writer.writerow(["CNPJ", dados_cnpj.get('cnpj', 'Não encontrado')])
        writer.writerow(["RAZÃO SOCIAL", dados_cnpj.get('nome', 'Não encontrado')])
        writer.writerow(["MATRIZ OU FILIAL", dados_cnpj.get('tipo', 'Não encontrado')])
        writer.writerow(["NOME FANTASIA", dados_cnpj.get('fantasia', 'Não encontrado')])
        writer.writerow(["SITUAÇÃO CADASTRAL", dados_cnpj.get('situacao', 'Não encontrado')])
        writer.writerow(["DATA DA SITUAÇÃO CADASTRAL", dados_cnpj.get('data_situacao', 'Não encontrado')])
        writer.writerow(["MOTIVO DA SITUAÇÃO CADASTRAL", dados_cnpj.get('motivo_situacao', 'Não encontrado')])
        writer.writerow(["NATUREZA JURÍDICA", dados_cnpj.get('natureza_juridica', 'Não encontrado')])
        writer.writerow(["DATA DE ABERTURA", dados_cnpj.get('abertura', 'Não encontrado')])
        writer.writerow(["IDADE", calcular_idade(dados_cnpj['abertura']) if 'abertura' in dados_cnpj else 'Não encontrado'])
        writer.writerow(["PORTE (RFB)", dados_cnpj.get('porte', 'Não encontrado')])
        writer.writerow(["CAPITAL SOCIAL", f"R$ {dados_cnpj.get('capital_social', 'Não encontrado')}"])
        writer.writerow(["ATUALIZAÇÃO DESTA PÁGINA", dados_cnpj.get('ultima_atualizacao', 'Não encontrado')])
        writer.writerow(["", ""])

        writer.writerow(["\nLocalização\n", ""])

        writer.writerow(["ENDEREÇO", f"{dados_cnpj.get('logradouro', 'Não encontrado')} | Número: {dados_cnpj.get('numero', 'Não encontrado')}"])
        writer.writerow(["COMPLEMENTO", dados_cnpj.get('complemento', 'Não encontrado')])
        writer.writerow(["BAIRRO", dados_cnpj.get('bairro', 'Não encontrado')])
        writer.writerow(["CIDADE | ESTADO", f"{dados_cnpj.get('municipio', 'Não encontrado')} | {dados_cnpj.get('uf', 'Não encontrado')}"])
        writer.writerow(["CEP", dados_cnpj.get('cep', 'Não encontrado')])
        writer.writerow(["TELEFONES", dados_cnpj.get('telefone', 'Não encontrado')])
        writer.writerow(["E-MAILS", dados_cnpj.get('email', 'Não encontrado')])
        writer.writerow(["", ""])

        writer.writerow(["Atividade Econômica Principal", ""])

        writer.writerow(["CÓDIGO", dados_cnpj['atividade_principal'][0]['code'] if 'atividade_principal' in dados_cnpj and dados_cnpj['atividade_principal'] else 'Não encontrado'])
        writer.writerow(["DESCRIÇÃO", dados_cnpj['atividade_principal'][0]['text'] if 'atividade_principal' in dados_cnpj and dados_cnpj['atividade_principal'] else 'Não encontrado'])
        writer.writerow(["", ""])
        
        writer.writerow(["Atividades Econômicas Secundárias", ""])
        if 'atividades_secundarias' in dados_cnpj and dados_cnpj['atividades_secundarias']:
            for atividade in dados_cnpj['atividades_secundarias']:
                writer.writerow(["CÓDIGO | DESCRIÇÃO", f"{atividade['code']} | {atividade['text']}"])
        else:
            writer.writerow(["", "Não encontrado"])
        writer.writerow(["", ""])
        writer.writerow(["Quadro de Sócios e Administradores (QSA)", ""])
        if 'qsa' in dados_cnpj and dados_cnpj['qsa']:
            for socio in dados_cnpj['qsa']:
                writer.writerow(["NOME", socio.get('nome', 'Não encontrado')])
                writer.writerow(["QUALIFICAÇÃO", socio.get('qual', 'Não encontrado')])
                data_entrada = socio.get('data_entrada', None)
                if data_entrada:
                    try:
                        data_entrada = datetime.strptime(data_entrada, '%d/%m/%Y').strftime('%d/%m/%Y')
                        writer.writerow(["ENTRADA", data_entrada])
                    except ValueError:
                        writer.writerow(["ENTRADA", "Data inválida"])
        else:
            writer.writerow(["", "Não encontrado"])

def salvar_json(dados_cnpj, filename):
    with open(filename, 'w', encoding='utf-8') as f:
        json.dump(dados_cnpj, f, indent=4, ensure_ascii=False)

def salvar_html(dados_cnpj, filename):
    html_content = f"""
    <!DOCTYPE html>
    <html lang="pt-BR">
    <head>
        <meta charset="UTF-8">
        <title>Consulta CNPJ</title>
        <style>
            body {{ font-family: Arial, sans-serif; margin: 20px; }}
            h1 {{ text-align: center; }}
            h2 {{ color: #333; }}
            p {{ margin: 5px 0; }}
        </style>
    </head>
    <body>
        <h1>Consulta CNPJ</h1>
        <h2>Informações Gerais</h2>
        <p><b>CNPJ:</b> {dados_cnpj.get('cnpj', 'Não encontrado')}</p>
        <p><b>RAZÃO SOCIAL:</b> {dados_cnpj.get('nome', 'Não encontrado')}</p>
        <p><b>MATRIZ OU FILIAL:</b> {dados_cnpj.get('tipo', 'Não encontrado')}</p>
        <p><b>NOME FANTASIA:</b> {dados_cnpj.get('fantasia', 'Não encontrado')}</p>
        <p><b>SITUAÇÃO CADASTRAL:</b> {dados_cnpj.get('situacao', 'Não encontrado')}</p>
        <p><b>DATA DA SITUAÇÃO CADASTRAL:</b> {dados_cnpj.get('data_situacao', 'Não encontrado')}</p>
        <p><b>MOTIVO DA SITUAÇÃO CADASTRAL:</b> {dados_cnpj.get('motivo_situacao', 'Não encontrado')}</p>
        <p><b>NATUREZA JURÍDICA:</b> {dados_cnpj.get('natureza_juridica', 'Não encontrado')}</p>
        <p><b>DATA DE ABERTURA:</b> {dados_cnpj.get('abertura', 'Não encontrado')}</p>
        <p><b>IDADE:</b> {calcular_idade(dados_cnpj['abertura']) if 'abertura' in dados_cnpj else 'Não encontrado'}</p>
        <p><b>PORTE (RFB):</b> {dados_cnpj.get('porte', 'Não encontrado')}</p>
        <p><b>CAPITAL SOCIAL:</b> R$ {dados_cnpj.get('capital_social', 'Não encontrado')}</p>
        <p><b>ATUALIZAÇÃO DESTA PÁGINA:</b> {dados_cnpj.get('ultima_atualizacao', 'Não encontrado')}</p>
        <h2>Localização</h2>
        <p><b>ENDEREÇO:</b> {dados_cnpj.get('logradouro', 'Não encontrado')} | Número: {dados_cnpj.get('numero', 'Não encontrado')}</p>
        <p><b>COMPLEMENTO:</b> {dados_cnpj.get('complemento', 'Não encontrado')}</p>
        <p><b>BAIRRO:</b> {dados_cnpj.get('bairro', 'Não encontrado')}</p>
        <p><b>CIDADE | ESTADO:</b> {dados_cnpj.get('municipio', 'Não encontrado')} | {dados_cnpj.get('uf', 'Não encontrado')}</p>
        <p><b>CEP:</b> {dados_cnpj.get('cep', 'Não encontrado')}</p>
        <p><b>TELEFONES:</b> {dados_cnpj.get('telefone', 'Não encontrado')}</p>
        <p><b>E-MAILS:</b> {dados_cnpj.get('email', 'Não encontrado')}</p>
        <h2>Atividade Econômica Principal</h2>
        <p><b>CÓDIGO:</b> {dados_cnpj['atividade_principal'][0]['code'] if 'atividade_principal' in dados_cnpj and dados_cnpj['atividade_principal'] else 'Não encontrado'}</p>
        <p><b>DESCRIÇÃO:</b> {dados_cnpj['atividade_principal'][0]['text'] if 'atividade_principal' in dados_cnpj and dados_cnpj['atividade_principal'] else 'Não encontrado'}</p>
        <h2>Atividades Econômicas Secundárias</h2>
    """
    if 'atividades_secundarias' in dados_cnpj and dados_cnpj['atividades_secundarias']:
        for atividade in dados_cnpj['atividades_secundarias']:
            html_content += f"<p><b>CÓDIGO:</b> {atividade['code']} | <b>DESCRIÇÃO:</b> {atividade['text']}</p>\n"
    else:
        html_content += "<p>Não encontrado</p>\n"

    html_content += """
        <h2>Quadro de Sócios e Administradores (QSA)</h2>
    """
    if 'qsa' in dados_cnpj and dados_cnpj['qsa']:
        for socio in dados_cnpj['qsa']:
            html_content += f"<p><b>NOME:</b> {socio.get('nome', 'Não encontrado')}</p>\n"
            html_content += f"<p><b>QUALIFICAÇÃO:</b> {socio.get('qual', 'Não encontrado')}</p>\n"
            data_entrada = socio.get('data_entrada', None)
            if data_entrada:
                try:
                    data_entrada = datetime.strptime(data_entrada, '%d/%m/%Y').strftime('%d/%m/%Y')
                    html_content += f"<p><b>ENTRADA:</b> {data_entrada}</p>\n"
                except ValueError:
                    html_content += "<p><b>ENTRADA:</b> Data inválida</p>\n"
    else:
        html_content += "<p>Não encontrado</p>\n"

    html_content += """
    </body>
    </html>
    """
    with open(filename, 'w', encoding='utf-8') as f:
        f.write(html_content)

def salvar_xml(dados_cnpj, filename):
    root = ET.Element("ConsultaCNPJ")
    
    geral = ET.SubElement(root, "InformacoesGerais")
    ET.SubElement(geral, "CNPJ").text = dados_cnpj.get('cnpj', 'Não encontrado')
    ET.SubElement(geral, "RazaoSocial").text = dados_cnpj.get('nome', 'Não encontrado')
    ET.SubElement(geral, "MatrizOuFilial").text = dados_cnpj.get('tipo', 'Não encontrado')
    ET.SubElement(geral, "NomeFantasia").text = dados_cnpj.get('fantasia', 'Não encontrado')
    ET.SubElement(geral, "SituacaoCadastral").text = dados_cnpj.get('situacao', 'Não encontrado')
    ET.SubElement(geral, "DataSituacaoCadastral").text = dados_cnpj.get('data_situacao', 'Não encontrado')
    ET.SubElement(geral, "MotivoSituacaoCadastral").text = dados_cnpj.get('motivo_situacao', 'Não encontrado')
    ET.SubElement(geral, "NaturezaJuridica").text = dados_cnpj.get('natureza_juridica', 'Não encontrado')
    ET.SubElement(geral, "DataAbertura").text = dados_cnpj.get('abertura', 'Não encontrado')
    ET.SubElement(geral, "Idade").text = calcular_idade(dados_cnpj['abertura']) if 'abertura' in dados_cnpj else 'Não encontrado'
    ET.SubElement(geral, "PorteRFB").text = dados_cnpj.get('porte', 'Não encontrado')
    ET.SubElement(geral, "CapitalSocial").text = f"R$ {dados_cnpj.get('capital_social', 'Não encontrado')}"
    ET.SubElement(geral, "UltimaAtualizacao").text = dados_cnpj.get('ultima_atualizacao', 'Não encontrado')

    localizacao = ET.SubElement(root, "Localizacao")
    ET.SubElement(localizacao, "Endereco").text = f"{dados_cnpj.get('logradouro', 'Não encontrado')} | Número: {dados_cnpj.get('numero', 'Não encontrado')}"
    ET.SubElement(localizacao, "Complemento").text = dados_cnpj.get('complemento', 'Não encontrado')
    ET.SubElement(localizacao, "Bairro").text = dados_cnpj.get('bairro', 'Não encontrado')
    ET.SubElement(localizacao, "CidadeEstado").text = f"{dados_cnpj.get('municipio', 'Não encontrado')} | {dados_cnpj.get('uf', 'Não encontrado')}"
    ET.SubElement(localizacao, "CEP").text = dados_cnpj.get('cep', 'Não encontrado')
    ET.SubElement(localizacao, "Telefones").text = dados_cnpj.get('telefone', 'Não encontrado')
    ET.SubElement(localizacao, "Emails").text = dados_cnpj.get('email', 'Não encontrado')

    atividade_principal = ET.SubElement(root, "AtividadeEconomicaPrincipal")
    ET.SubElement(atividade_principal, "Codigo").text = dados_cnpj['atividade_principal'][0]['code'] if 'atividade_principal' in dados_cnpj and dados_cnpj['atividade_principal'] else 'Não encontrado'
    ET.SubElement(atividade_principal, "Descricao").text = dados_cnpj['atividade_principal'][0]['text'] if 'atividade_principal' in dados_cnpj and dados_cnpj['atividade_principal'] else 'Não encontrado'

    atividades_secundarias = ET.SubElement(root, "AtividadesEconomicasSecundarias")
    if 'atividades_secundarias' in dados_cnpj and dados_cnpj['atividades_secundarias']:
        for atividade in dados_cnpj['atividades_secundarias']:
            atividade_elem = ET.SubElement(atividades_secundarias, "Atividade")
            ET.SubElement(atividade_elem, "Codigo").text = atividade['code']
            ET.SubElement(atividade_elem, "Descricao").text = atividade['text']
    else:
        ET.SubElement(atividades_secundarias, "Info").text = "Não encontrado"

    qsa = ET.SubElement(root, "QuadroSociosAdministradores")
    if 'qsa' in dados_cnpj and dados_cnpj['qsa']:
        for socio in dados_cnpj['qsa']:
            socio_elem = ET.SubElement(qsa, "Socio")
            ET.SubElement(socio_elem, "Nome").text = socio.get('nome', 'Não encontrado')
            ET.SubElement(socio_elem, "Qualificacao").text = socio.get('qual', 'Não encontrado')
            data_entrada = socio.get('data_entrada', None)
            if data_entrada:
                try:
                    data_entrada = datetime.strptime(data_entrada, '%d/%m/%Y').strftime('%d/%m/%Y')
                    ET.SubElement(socio_elem, "Entrada").text = data_entrada
                except ValueError:
                    ET.SubElement(socio_elem, "Entrada").text = "Data inválida"
    else:
        ET.SubElement(qsa, "Info").text = "Não encontrado"

    tree = ET.ElementTree(root)
    tree.write(filename)

def salvar_arquivo(dados_cnpj, formato_var):
    formato = formato_var.get()
    cnpj = dados_cnpj.get('cnpj', 'desconhecido').replace('/', '_').replace('.', '_').replace('-', '_')
    
    format_configs = {
        "Excel (.xlsx)": (".xlsx", [("Excel files", "*.xlsx")], f"consulta_cnpj_{cnpj}.xlsx", salvar_xlsx),
        "Word (.docx)": (".docx", [("Word files", "*.docx")], f"consulta_cnpj_{cnpj}.docx", salvar_docx),
        "Texto (.txt)": (".txt", [("Text files", "*.txt")], f"consulta_cnpj_{cnpj}.txt", salvar_txt),
        "CSV (.csv)": (".csv", [("CSV files", "*.csv")], f"consulta_cnpj_{cnpj}.csv", salvar_csv),
        "JSON (.json)": (".json", [("JSON files", "*.json")], f"consulta_cnpj_{cnpj}.json", salvar_json),
        "HTML (.html)": (".html", [("HTML files", "*.html")], f"consulta_cnpj_{cnpj}.html", salvar_html),
        "XML (.xml)": (".xml", [("XML files", "*.xml")], f"consulta_cnpj_{cnpj}.xml", salvar_xml)
    }

    if formato not in format_configs:
        messagebox.showerror("Erro", "Formato de arquivo inválido")
        return

    ext, filetypes, default_name, save_func = format_configs[formato]
    filename = filedialog.asksaveasfilename(
        title="Salvar como",
        defaultextension=ext,
        filetypes=filetypes,
        initialfile=default_name
    )
    
    if filename:
        try:
            save_func(dados_cnpj, filename)
            messagebox.showinfo("Sucesso", f"Arquivo Salvo\n\n{filename}")
        except Exception as e:
            messagebox.showerror("Erro ao salvar", f"Não foi possível salvar o arquivo: {str(e)}")

def consultar_e_mostrar(cnpj_entry, info_text, maps_button, salvar_button, formato_var):
    cnpj = limpar_cnpj(cnpj_entry.get())
    if not cnpj:
        messagebox.showerror("Erro", "Por favor, insira um CNPJ válido")
        return
    dados_cnpj = consultar_cnpj(cnpj)
    if dados_cnpj:
        info_text.config(state=tk.NORMAL)
        info_text.delete(1.0, tk.END)
        
        message = gerar_texto(dados_cnpj)
        info_text.insert(tk.END, message)
        info_text.config(state=tk.DISABLED)

        logradouro = dados_cnpj.get('logradouro', 'Não encontrado')
        numero = dados_cnpj.get('numero', 'Não encontrado')
        municipio = dados_cnpj.get('municipio', 'Não encontrado')
        uf = dados_cnpj.get('uf', 'Não encontrado')

        maps_button.config(state=tk.NORMAL, command=lambda: abrir_google_maps(logradouro, numero, municipio, uf))
        salvar_button.config(state=tk.NORMAL, command=lambda: salvar_arquivo(dados_cnpj, formato_var))

# ==================== INTERFACE ====================

def criar_interface_grafica():
    global text_area
    root = tk.Tk()
    root.title("CNPJ CONSULTAR")
    root.geometry("1280x900")
    root.state('zoomed')
    root.configure(bg=HackerStyle.BG)

    # Título Cyber
    title = tk.Label(root, text="CNPJ CONSULTAR",
                     font=HackerStyle.TITLE_FONT,
                     fg=HackerStyle.FG, bg=HackerStyle.BG)
    title.pack(pady=15)

    # ==================== FRAME PRINCIPAL (TODOS OS BOTÕES NA MESMA LINHA) ====================
    top_frame = tk.Frame(root, bg=HackerStyle.BG)
    top_frame.pack(pady=10, fill="x", padx=40)

    # CNPJ Label + Entry
    tk.Label(top_frame, text="CNPJ →", font=HackerStyle.FONT_BOLD,
             fg=HackerStyle.FG, bg=HackerStyle.BG).pack(side="left", padx=(0, 8))

    cnpj_entry = tk.Entry(top_frame, width=32, font=("Consolas", 14),
                          bg="#111111", fg=HackerStyle.FG, insertbackground=HackerStyle.FG)
    cnpj_entry.pack(side="left", padx=4)

    # Botão Consultar
    consultar_button = tk.Button(top_frame, text="▶ CONSULTAR", font=HackerStyle.FONT_BOLD,
                                 bg="#00ff41", fg="black", activebackground="#00cc33", width=12)
    consultar_button.pack(side="left", padx=4)

    # Botão Google Maps
    maps_button = tk.Button(top_frame, text="🗺️GOOGLE MAPS", font=HackerStyle.FONT,
                            bg="#00FFFF", fg="black", state=tk.DISABLED, width=16)
    maps_button.pack(side="left", padx=4)

    # Botão Salvar (agora na mesma linha)
    salvar_button = tk.Button(top_frame, text="💾SALVAR RELATÓRIO", font=HackerStyle.FONT_BOLD,
                              bg="#FF4500", fg="black", state=tk.DISABLED, width=19)
    salvar_button.pack(side="left", padx=4)

    # Botão Abrir Arquivo (agora na mesma linha)
    btn_abrir = tk.Button(top_frame, text="📂ABRIR ARQUIVO", command=abrir_arquivo,
                          font=HackerStyle.FONT_BOLD, bg="#00ff41", fg="black", width=16)
    btn_abrir.pack(side="left", padx=4)

    # Menu de Formato (à direita)
    formato_var = tk.StringVar(root)
    formato_var.set("Excel (.xlsx)")
    formato_menu = tk.OptionMenu(top_frame, formato_var,
                                 "Excel (.xlsx)", "Word (.docx)", "Texto (.txt)",
                                 "CSV (.csv)", "JSON (.json)", "HTML (.html)", "XML (.xml)")
    formato_menu.config(font=HackerStyle.FONT, bg=HackerStyle.GRAY, fg=HackerStyle.FG, width=100)
    formato_menu.pack(side="right", padx=2)

    # ==================== ÁREA DE TEXTO (TERMINAL) ====================
    text_area = scrolledtext.ScrolledText(
        root, wrap=tk.WORD, width=145, height=38,
        font=("Consolas", 10), bg="#000000", fg=HackerStyle.FG,
        insertbackground=HackerStyle.FG, selectbackground="#00ff41", selectforeground="black"
    )
    text_area.pack(pady=20, padx=40, fill="both", expand=True)
    text_area.config(state=tk.DISABLED)

    # ==================== CONECTAR COMANDOS ====================
    consultar_button.config(command=lambda: consultar_e_mostrar(
        cnpj_entry, text_area, maps_button, salvar_button, formato_var))

    root.mainloop()

if __name__ == "__main__":    
    criar_interface_grafica()
