import socket
import re
import requests
from datetime import datetime
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from tkinter.scrolledtext import ScrolledText
import webbrowser
import os
import json
import csv
import xml.etree.ElementTree as ET
from xml.dom import minidom
import openpyxl
from openpyxl.utils import get_column_letter
from docx import Document

# ==================== ESTILO ====================
class HackerStyle:
    BG = "#0a0a0a"
    FG = "#00ff41"
    ACCENT = "#00ff41"
    RED = "#ff0033"
    GRAY = "#1a1a1a"
    FONT = ("Consolas", 11)
    FONT_BOLD = ("Consolas", 12, "bold")
    TITLE_FONT = ("Consolas", 18, "bold")

def abrir_arquivo():
    caminho = filedialog.askopenfilename(
        title="Abrir arquivo salvo",
        filetypes=[
            ("Todos os suportados", "*.txt *.docx *.xlsx *.csv *.json *.log *.md *.html *.htm *.xml"),
            ("Texto", "*.txt *.log *.md"),
            ("Word", "*.docx"),
            ("Excel", "*.xlsx"),
            ("CSV", "*.csv"),
            ("JSON", "*.json"),
            ("HTML", "*.html *.htm"),
            ("XML", "*.xml"),
            ("Todos os arquivos", "*.*")
        ]
    )
    if not caminho:
        return

    try:
        ext = os.path.splitext(caminho)[1].lower()
        conteudo = ""

        if ext in [".txt", ".log", ".md"]:
            with open(caminho, "r", encoding="utf-8") as f:
                conteudo = f.read()

        elif ext in [".html", ".htm"]:
            with open(caminho, "r", encoding="utf-8") as f:
                conteudo = f.read()

        elif ext == ".xml":
            try:
                tree = ET.parse(caminho)
                conteudo = minidom.parseString(ET.tostring(tree.getroot(), encoding='unicode')).toprettyxml(indent="  ")
            except:
                with open(caminho, "r", encoding="utf-8") as f:
                    conteudo = f.read()

        elif ext == ".json":
            with open(caminho, "r", encoding="utf-8") as f:
                try:
                    dados = json.load(f)
                    conteudo = json.dumps(dados, indent=4, ensure_ascii=False)
                except:
                    conteudo = "Erro: Arquivo JSON malformado"

        elif ext == ".csv":
            with open(caminho, "r", encoding="utf-8") as f:
                conteudo = f.read()

        elif ext == ".docx":
            try:
                doc = Document(caminho)
                conteudo = "\n\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            except Exception as e:
                conteudo = f"Erro ao ler DOCX: {str(e)}"

        elif ext == ".xlsx":
            try:
                wb = openpyxl.load_workbook(caminho, data_only=True)
                sheet = wb.active
                linhas = []
                for row in sheet.iter_rows(values_only=True):
                    linha = "\t".join([str(cell) if cell is not None else "" for cell in row])
                    linhas.append(linha)
                conteudo = "\n".join(linhas)
            except Exception as e:
                conteudo = f"Erro ao ler XLSX: {str(e)}"

        else:
            messagebox.showerror("Erro", f"Tipo de arquivo não suportado: {ext}")
            return

        # ==================== MOSTRAR NO TEXT AREA ====================
        # Escolhe o text_area mais apropriado (prioridade para aba atual)
        if notebook.select() == str(aba_cnpj):
            widget = text_cnpj
        elif notebook.select() == str(aba_whois):
            widget = text_whois
        else:
            widget = text_cnpj  # default

        widget.config(state=tk.NORMAL)
        widget.delete(1.0, tk.END)
        widget.insert(tk.END, f"--- ARQUIVO ABERTO: {os.path.basename(caminho)} ---\n\n")
        widget.insert(tk.END, conteudo)
        widget.config(state=tk.DISABLED)

        messagebox.showinfo("Sucesso", f"Arquivo aberto com sucesso!\n\n{os.path.basename(caminho)}")

    except Exception as e:
        messagebox.showerror("Erro ao abrir arquivo", f"Erro: {str(e)}")

# ===================== TRADUÇÃO WHOIS =====================
traducao = {
    "domain:": "Domínio",
    "owner:": "Entidade",
    "ownerid:": "CNPJ",
    "responsible:": "Responsável",
    "country:": "País",
    "created:": "Criado em",
    "changed:": "Alterado em",
    "expires:": "Expira em",
    "status:": "Status",
    "nserver:": "Servidor DNS",
    "nameserver:": "Servidor DNS",
    "nameservers:": "Servidores DNS",
    "person:": "Pessoa",
    "e-mail:": "E-mail",
    "email:": "E-mail",
    "inetnum:": "Faixa de IP",
    "netname:": "Nome da Rede",
    "descr:": "Descrição",
    "org:": "Organização",
    "address:": "Endereço",
    "phone:": "Telefone",
    "abuse-mailbox:": "E-mail de Abuso",
    "source:": "Fonte",
    # Campos .gov e internacionais
    "registrar:": "Registrador",
    "registrant:": "Registrante",
    "registrant organization:": "Organização Registrante",
    "registrant street:": "Endereço",
    "registrant city:": "Cidade",
    "registrant state/province:": "Estado/Província",
    "registrant postal code:": "CEP",
    "registrant country:": "País",
    "registrant phone:": "Telefone",
    "registrant email:": "E-mail",
    "admin:": "Administrador",
    "tech:": "Técnico",
    "name server:": "Servidor DNS",
    "dnssec:": "DNSSEC",
    "domain status:": "Status do Domínio",
    "updated date:": "Atualizado em",
    "creation date:": "Criado em",
    "registry expiry date:": "Expira em",
}

def formatar_data_brasileira(texto):
    formatos = ["%Y-%m-%d", "%Y%m%d", "%Y-%m-%dT%H:%M:%SZ", "%Y-%m-%dT%H:%M:%S.%fZ"]
    for formato in formatos:
        try:
            data = datetime.strptime(texto.strip(), formato)
            return data.strftime("%d/%m/%Y")
        except:
            continue
    return texto

def traduzir_linha(linha):
    linha_lower = linha.lower()
    for termo, traducao_pt in traducao.items():
        if linha_lower.startswith(termo):
            valor = linha[len(termo):].strip()
            return f"{traducao_pt:<42}: {valor}"
    if ":" in linha:
        campo, valor = linha.split(":", 1)
        campo = campo.strip()
        return f"{campo:<42}: {valor.strip()}"
    return linha

def consultar_whois(entrada):
    try:
        # Detecta tipo
        try:
            socket.inet_pton(socket.AF_INET, entrada)
            tipo = "ipv4"
        except:
            try:
                socket.inet_pton(socket.AF_INET6, entrada)
                tipo = "ipv6"
            except:
                tipo = "dominio"

        if tipo in ["ipv4", "ipv6"]:
            servidor = 'whois.iana.org'
            with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
                s.settimeout(10)
                s.connect((servidor, 43))
                s.send((entrada + "\r\n").encode())
                resposta = b""
                while True:
                    dados = s.recv(4096)
                    if not dados: break
                    resposta += dados
            texto_iana = resposta.decode(errors='ignore')
            match = re.search(r"refer:\s*(\S+)", texto_iana, re.IGNORECASE)
            servidor = match.group(1) if match else 'whois.arin.net'
        else:
            tld = '.' + entrada.split('.')[-1].lower()
            servidores_whois_tld = {
                '.com': 'whois.verisign-grs.com',
                '.net': 'whois.verisign-grs.com',
                '.org': 'whois.pir.org',
                '.br': 'whois.registro.br',
                '.gov': 'whois.nic.gov',
                '.edu': 'whois.educause.edu',
            }
            servidor = servidores_whois_tld.get(tld)
            if not servidor:
                return "TLD não suportado no momento."

        # Consulta WHOIS
        with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
            s.settimeout(15)
            s.connect((servidor, 43))
            s.send((entrada + "\r\n").encode())
            resposta = b""
            while True:
                dados = s.recv(4096)
                if not dados: break
                resposta += dados

        texto = resposta.decode(errors='ignore')

        # ==================== LIMPEZA DE DISCLAIMERS ====================
        texto = re.sub(
            r'(Information.*?support.*?access.*?)(\n\n|\Z)',
            '',
            texto,
            flags=re.IGNORECASE | re.DOTALL
        )

        linhas = texto.splitlines()
        saida_formatada = ["=" * 90, f"WHOIS → {entrada.upper()}", "=" * 90, ""]

        for linha in linhas:
            linha = linha.strip()
            if not linha:
                continue

            linha_lower = linha.lower()
            
            if re.search(r'copyright|terms|usage|legal|reserved|icann|verisign|notice|for more|information is provided', linha_lower):
                continue

            disclaimers = ["information is provided", "informational purposes", "as is without", "guarantee of accuracy"]
            if any(frase in linha_lower for frase in disclaimers):
                continue

            if linha.startswith(('%', '#', '>>>', '---', '==')):
                continue

            # Formatar datas
            linha = re.sub(r"\d{4}-\d{2}-\d{2}(T[\d:.Z]+)?|\d{8}",
                          lambda m: formatar_data_brasileira(m.group()), linha)

            linha_traduzida = traduzir_linha(linha)
            saida_formatada.append(linha_traduzida)

        return "\n".join(saida_formatada)

    except Exception as e:
        return f"[-] Erro na consulta: {e}"
      

# ===================== CONSULTA CNPJ AVANÇADA =====================
def limpar_cnpj(cnpj):
    return ''.join(filter(str.isdigit, cnpj))

def consultar_cnpj(cnpj):
    cnpj = limpar_cnpj(cnpj)
    if len(cnpj) != 14:
        return None
    try:
        r = requests.get(f"https://www.receitaws.com.br/v1/cnpj/{cnpj}", timeout=12)
        if r.status_code == 200:
            return r.json()
        return None
    except:
        return None

def calcular_idade(data_abertura):
    try:
        hoje = datetime.now()
        data = datetime.strptime(data_abertura, '%d/%m/%Y')
        diferenca = hoje - data
        anos = diferenca.days // 365
        meses = (diferenca.days % 365) // 30
        dias = (diferenca.days % 365) % 30
        return f"{anos} anos, {meses} meses e {dias} dias"
    except:
        return "Data inválida"

def abrir_google_maps(logradouro, numero, municipio, uf):
    endereco = f"{logradouro}, {numero}, {municipio}, {uf}"
    url = f"https://www.google.com/maps/search/?api=1&query={endereco.replace(' ', '+')}"
    webbrowser.open(url)

def limpar_cnpj(cnpj):
    return ''.join(filter(str.isdigit, cnpj))

def gerar_texto(dados_cnpj):
    logradouro = dados_cnpj.get('logradouro', 'Não encontrado')
    numero = dados_cnpj.get('numero', 'Não encontrado')
    municipio = dados_cnpj.get('municipio', 'Não encontrado')
    uf = dados_cnpj.get('uf', 'Não encontrado')

    message = f"""
CNPJ: {dados_cnpj.get('cnpj', 'Não encontrado')}

RAZÃO SOCIAL: {dados_cnpj.get('nome', 'Não encontrado')}

MATRIZ OU FILIAL: {dados_cnpj.get('tipo', 'Não encontrado')}

NOME FANTASIA: {dados_cnpj.get('fantasia', 'Não encontrado')}

SITUAÇÃO CADASTRAL: {dados_cnpj.get('situacao', 'Não encontrado')}

DATA DA SITUAÇÃO CADASTRAL: {dados_cnpj.get('data_situacao', 'Não encontrado')}

MOTIVO DA SITUAÇÃO CADASTRAL: {dados_cnpj.get('motivo_situacao', 'Não encontrado')}

NATUREZA JURÍDICA: {dados_cnpj.get('natureza_juridica', 'Não encontrado')}

DATA DE ABERTURA: {dados_cnpj.get('abertura', 'Não encontrado')}

IDADE: {calcular_idade(dados_cnpj['abertura']) if 'abertura' in dados_cnpj else 'Não encontrado'}

PORTE (RFB): {dados_cnpj.get('porte', 'Não encontrado')}

CAPITAL SOCIAL: R$ {dados_cnpj.get('capital_social', 'Não encontrado')}

ATUALIZAÇÃO DESTA PÁGINA: {dados_cnpj.get('ultima_atualizacao', 'Não encontrado')}



LOCALIZAÇÃO
===========

ENDEREÇO: {logradouro}       |  Número: {numero}

COMPLEMENTO: {dados_cnpj.get('complemento', 'Não encontrado')}

BAIRRO: {dados_cnpj.get('bairro', 'Não encontrado')}

CIDADE | ESTADO: {municipio} | {uf}

CEP: {dados_cnpj.get('cep', 'Não encontrado')}

TELEFONES: {dados_cnpj.get('telefone', 'Não encontrado')}

E-MAILS: {dados_cnpj.get('email', 'Não encontrado')}



ATIVIDADE ECONÔMICA PRINCIPAL
==============================

CÓDIGO: {dados_cnpj['atividade_principal'][0]['code'] if 'atividade_principal' in dados_cnpj and dados_cnpj['atividade_principal'] else 'Não encontrado'}
DESCRIÇÃO: {dados_cnpj['atividade_principal'][0]['text'] if 'atividade_principal' in dados_cnpj and dados_cnpj['atividade_principal'] else 'Não encontrado'}


ATIVIDADES ECONÔMICAS SECUNDÁRIAS
=================================

"""
    if 'atividades_secundarias' in dados_cnpj and dados_cnpj['atividades_secundarias']:
        for atividade in dados_cnpj['atividades_secundarias']:
            message += f"CÓDIGO: {atividade['code']} | DESCRIÇÃO: {atividade['text']}\n"
    else:
        message += "Não encontrado\n"

    message += "\n\nQUADRO DE SÓCIOS E ADMINISTRADORES (QSA)\n==========================================\n"
    if 'qsa' in dados_cnpj and dados_cnpj['qsa']:
        for socio in dados_cnpj['qsa']:
            data_entrada = socio.get('data_entrada', None)
            if data_entrada:
                try:
                    data_entrada = datetime.strptime(data_entrada, '%d/%m/%Y').strftime('%d/%m/%Y')
                    message += f"""
NOME: {socio.get('nome', 'Não encontrado')}
QUALIFICAÇÃO: {socio.get('qual', 'Não encontrado')}
ENTRADA: {data_entrada}
"""
                except ValueError:
                    message += f"""
NOME: {socio.get('nome', 'Não encontrado')}
QUALIFICAÇÃO: {socio.get('qual', 'Não encontrado')}
ENTRADA: Data inválida
"""
            else:
                message += f"""
NOME: {socio.get('nome', 'Não encontrado')}
QUALIFICAÇÃO: {socio.get('qual', 'Não encontrado')}
"""
    else:
        message += "Não encontrado\n"
    
    return message

def salvar_xlsx(dados_cnpj, filename):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Consulta CNPJ"

    headers = ["Categoria", "Informação\n"]

    ws.append(headers)

    ws.append(["Informações Gerais\n", ""])

    ws.append(["CNPJ", dados_cnpj.get('cnpj', 'Não encontrado')])
    ws.append(["RAZÃO SOCIAL", dados_cnpj.get('nome', 'Não encontrado')])
    ws.append(["MATRIZ OU FILIAL", dados_cnpj.get('tipo', 'Não encontrado')])
    ws.append(["NOME FANTASIA", dados_cnpj.get('fantasia', 'Não encontrado')])
    ws.append(["SITUAÇÃO CADASTRAL", dados_cnpj.get('situacao', 'Não encontrado')])
    ws.append(["DATA DA SITUAÇÃO CADASTRAL", dados_cnpj.get('data_situacao', 'Não encontrado')])
    ws.append(["MOTIVO DA SITUAÇÃO CADASTRAL", dados_cnpj.get('motivo_situacao', 'Não encontrado')])
    ws.append(["NATUREZA JURÍDICA", dados_cnpj.get('natureza_juridica', 'Não encontrado')])
    ws.append(["DATA DE ABERTURA", dados_cnpj.get('abertura', 'Não encontrado')])
    ws.append(["IDADE", calcular_idade(dados_cnpj['abertura']) if 'abertura' in dados_cnpj else 'Não encontrado'])
    ws.append(["PORTE (RFB)", dados_cnpj.get('porte', 'Não encontrado')])
    ws.append(["CAPITAL SOCIAL", f"R$ {dados_cnpj.get('capital_social', 'Não encontrado')}"])
    ws.append(["ATUALIZAÇÃO DESTA PÁGINA", dados_cnpj.get('ultima_atualizacao', 'Não encontrado')])
    ws.append(["", ""])

    ws.append(["\nLocalização\n", ""])

    ws.append(["ENDEREÇO", f"{dados_cnpj.get('logradouro', 'Não encontrado')} | Número: {dados_cnpj.get('numero', 'Não encontrado')}"])
    ws.append(["COMPLEMENTO", dados_cnpj.get('complemento', 'Não encontrado')])
    ws.append(["BAIRRO", dados_cnpj.get('bairro', 'Não encontrado')])
    ws.append(["CIDADE | ESTADO", f"{dados_cnpj.get('municipio', 'Não encontrado')} | {dados_cnpj.get('uf', 'Não encontrado')}"])
    ws.append(["CEP", dados_cnpj.get('cep', 'Não encontrado')])
    ws.append(["TELEFONES", dados_cnpj.get('telefone', 'Não encontrado')])
    ws.append(["E-MAILS", dados_cnpj.get('email', 'Não encontrado')])
    ws.append(["", ""])

    ws.append(["Atividade Econômica Principal", ""])
    ws.append(["CÓDIGO", dados_cnpj['atividade_principal'][0]['code'] if 'atividade_principal' in dados_cnpj and dados_cnpj['atividade_principal'] else 'Não encontrado'])
    ws.append(["DESCRIÇÃO", dados_cnpj['atividade_principal'][0]['text'] if 'atividade_principal' in dados_cnpj and dados_cnpj['atividade_principal'] else 'Não encontrado'])
    ws.append(["", ""])

    ws.append(["Atividades Econômicas Secundárias", ""])
    if 'atividades_secundarias' in dados_cnpj and dados_cnpj['atividades_secundarias']:
        for atividade in dados_cnpj['atividades_secundarias']:
            ws.append(["CÓDIGO | DESCRIÇÃO", f"{atividade['code']} | {atividade['text']}"])
    else:
        ws.append(["", "Não encontrado"])

    ws.append(["", ""])
    ws.append(["Quadro de Sócios e Administradores (QSA)", ""])
    if 'qsa' in dados_cnpj and dados_cnpj['qsa']:
        for socio in dados_cnpj['qsa']:
            ws.append(["NOME", socio.get('nome', 'Não encontrado')])
            ws.append(["QUALIFICAÇÃO", socio.get('qual', 'Não encontrado')])
            data_entrada = socio.get('data_entrada', None)
            if data_entrada:
                try:
                    data_entrada = datetime.strptime(data_entrada, '%d/%m/%Y').strftime('%d/%m/%Y')
                    ws.append(["ENTRADA", data_entrada])
                except ValueError:
                    ws.append(["ENTRADA", "Data inválida"])
    else:
        ws.append(["", "Não encontrado"])

    for col in range(1, ws.max_column + 1):
        ws.column_dimensions[get_column_letter(col)].width = 50

    wb.save(filename)

def salvar_docx(dados_cnpj, filename):
    doc = Document()
    doc.add_heading('Consulta CNPJ', 0)

    doc.add_heading('Informações Gerais', level=1)

    doc.add_paragraph(f"CNPJ: {dados_cnpj.get('cnpj', 'Não encontrado')}")
    doc.add_paragraph(f"RAZÃO SOCIAL: {dados_cnpj.get('nome', 'Não encontrado')}")
    doc.add_paragraph(f"MATRIZ OU FILIAL: {dados_cnpj.get('tipo', 'Não encontrado')}")
    doc.add_paragraph(f"NOME FANTASIA: {dados_cnpj.get('fantasia', 'Não encontrado')}")
    doc.add_paragraph(f"SITUAÇÃO CADASTRAL: {dados_cnpj.get('situacao', 'Não encontrado')}")
    doc.add_paragraph(f"DATA DA SITUAÇÃO CADASTRAL: {dados_cnpj.get('data_situacao', 'Não encontrado')}")
    doc.add_paragraph(f"MOTIVO DA SITUAÇÃO CADASTRAL: {dados_cnpj.get('motivo_situacao', 'Não encontrado')}")
    doc.add_paragraph(f"NATUREZA JURÍDICA: {dados_cnpj.get('natureza_juridica', 'Não encontrado')}")
    doc.add_paragraph(f"DATA DE ABERTURA: {dados_cnpj.get('abertura', 'Não encontrado')}")
    doc.add_paragraph(f"IDADE: {calcular_idade(dados_cnpj['abertura']) if 'abertura' in dados_cnpj else 'Não encontrado'}")
    doc.add_paragraph(f"PORTE (RFB): {dados_cnpj.get('porte', 'Não encontrado')}")
    doc.add_paragraph(f"CAPITAL SOCIAL: R$ {dados_cnpj.get('capital_social', 'Não encontrado')}")
    doc.add_paragraph(f"ATUALIZAÇÃO DESTA PÁGINA: {dados_cnpj.get('ultima_atualizacao', 'Não encontrado')}")

    doc.add_heading('\nLocalização\n', level=1)

    doc.add_paragraph(f"ENDEREÇO: {dados_cnpj.get('logradouro', 'Não encontrado')} | Número: {dados_cnpj.get('numero', 'Não encontrado')}")
    doc.add_paragraph(f"COMPLEMENTO: {dados_cnpj.get('complemento', 'Não encontrado')}")
    doc.add_paragraph(f"BAIRRO: {dados_cnpj.get('bairro', 'Não encontrado')}")
    doc.add_paragraph(f"CIDADE | ESTADO: {dados_cnpj.get('municipio', 'Não encontrado')} | {dados_cnpj.get('uf', 'Não encontrado')}")
    doc.add_paragraph(f"CEP: {dados_cnpj.get('cep', 'Não encontrado')}")
    doc.add_paragraph(f"TELEFONES: {dados_cnpj.get('telefone', 'Não encontrado')}")
    doc.add_paragraph(f"E-MAILS: {dados_cnpj.get('email', 'Não encontrado')}")
    doc.add_heading('Atividade Econômica Principal', level=1)
    doc.add_paragraph(f"CÓDIGO: {dados_cnpj['atividade_principal'][0]['code'] if 'atividade_principal' in dados_cnpj and dados_cnpj['atividade_principal'] else 'Não encontrado'}")
    doc.add_paragraph(f"DESCRIÇÃO: {dados_cnpj['atividade_principal'][0]['text'] if 'atividade_principal' in dados_cnpj and dados_cnpj['atividade_principal'] else 'Não encontrado'}")

    doc.add_heading('Atividades Econômicas Secundárias', level=1)
    if 'atividades_secundarias' in dados_cnpj and dados_cnpj['atividades_secundarias']:
        for atividade in dados_cnpj['atividades_secundarias']:
            doc.add_paragraph(f"CÓDIGO: {atividade['code']} | DESCRIÇÃO: {atividade['text']}")
    else:
        doc.add_paragraph("Não encontrado")

    doc.add_heading('Quadro de Sócios e Administradores (QSA)', level=1)
    if 'qsa' in dados_cnpj and dados_cnpj['qsa']:
        for socio in dados_cnpj['qsa']:
            doc.add_paragraph(f"NOME: {socio.get('nome', 'Não encontrado')}")
            doc.add_paragraph(f"QUALIFICAÇÃO: {socio.get('qual', 'Não encontrado')}")
            data_entrada = socio.get('data_entrada', None)
            if data_entrada:
                try:
                    data_entrada = datetime.strptime(data_entrada, '%d/%m/%Y').strftime('%d/%m/%Y')
                    doc.add_paragraph(f"ENTRADA: {data_entrada}")
                except ValueError:
                    doc.add_paragraph("ENTRADA: Data inválida")
    else:
        doc.add_paragraph("Não encontrado")

    doc.save(filename)

def salvar_txt(dados_cnpj, filename):
    with open(filename, 'w', encoding='utf-8') as f:
        f.write(gerar_texto(dados_cnpj))

def salvar_csv(dados_cnpj, filename):
    with open(filename, 'w', encoding='utf-8', newline='') as f:
        writer = csv.writer(f)
        writer.writerow(["Categoria", "Informação\n"])

        writer.writerow(["Informações Gerais\n", ""])

        writer.writerow(["CNPJ", dados_cnpj.get('cnpj', 'Não encontrado')])
        writer.writerow(["RAZÃO SOCIAL", dados_cnpj.get('nome', 'Não encontrado')])
        writer.writerow(["MATRIZ OU FILIAL", dados_cnpj.get('tipo', 'Não encontrado')])
        writer.writerow(["NOME FANTASIA", dados_cnpj.get('fantasia', 'Não encontrado')])
        writer.writerow(["SITUAÇÃO CADASTRAL", dados_cnpj.get('situacao', 'Não encontrado')])
        writer.writerow(["DATA DA SITUAÇÃO CADASTRAL", dados_cnpj.get('data_situacao', 'Não encontrado')])
        writer.writerow(["MOTIVO DA SITUAÇÃO CADASTRAL", dados_cnpj.get('motivo_situacao', 'Não encontrado')])
        writer.writerow(["NATUREZA JURÍDICA", dados_cnpj.get('natureza_juridica', 'Não encontrado')])
        writer.writerow(["DATA DE ABERTURA", dados_cnpj.get('abertura', 'Não encontrado')])
        writer.writerow(["IDADE", calcular_idade(dados_cnpj['abertura']) if 'abertura' in dados_cnpj else 'Não encontrado'])
        writer.writerow(["PORTE (RFB)", dados_cnpj.get('porte', 'Não encontrado')])
        writer.writerow(["CAPITAL SOCIAL", f"R$ {dados_cnpj.get('capital_social', 'Não encontrado')}"])
        writer.writerow(["ATUALIZAÇÃO DESTA PÁGINA", dados_cnpj.get('ultima_atualizacao', 'Não encontrado')])
        writer.writerow(["", ""])

        writer.writerow(["\nLocalização\n", ""])

        writer.writerow(["ENDEREÇO", f"{dados_cnpj.get('logradouro', 'Não encontrado')} | Número: {dados_cnpj.get('numero', 'Não encontrado')}"])
        writer.writerow(["COMPLEMENTO", dados_cnpj.get('complemento', 'Não encontrado')])
        writer.writerow(["BAIRRO", dados_cnpj.get('bairro', 'Não encontrado')])
        writer.writerow(["CIDADE | ESTADO", f"{dados_cnpj.get('municipio', 'Não encontrado')} | {dados_cnpj.get('uf', 'Não encontrado')}"])
        writer.writerow(["CEP", dados_cnpj.get('cep', 'Não encontrado')])
        writer.writerow(["TELEFONES", dados_cnpj.get('telefone', 'Não encontrado')])
        writer.writerow(["E-MAILS", dados_cnpj.get('email', 'Não encontrado')])
        writer.writerow(["", ""])

        writer.writerow(["Atividade Econômica Principal", ""])

        writer.writerow(["CÓDIGO", dados_cnpj['atividade_principal'][0]['code'] if 'atividade_principal' in dados_cnpj and dados_cnpj['atividade_principal'] else 'Não encontrado'])
        writer.writerow(["DESCRIÇÃO", dados_cnpj['atividade_principal'][0]['text'] if 'atividade_principal' in dados_cnpj and dados_cnpj['atividade_principal'] else 'Não encontrado'])
        writer.writerow(["", ""])
        
        writer.writerow(["Atividades Econômicas Secundárias", ""])
        if 'atividades_secundarias' in dados_cnpj and dados_cnpj['atividades_secundarias']:
            for atividade in dados_cnpj['atividades_secundarias']:
                writer.writerow(["CÓDIGO | DESCRIÇÃO", f"{atividade['code']} | {atividade['text']}"])
        else:
            writer.writerow(["", "Não encontrado"])
        writer.writerow(["", ""])
        writer.writerow(["Quadro de Sócios e Administradores (QSA)", ""])
        if 'qsa' in dados_cnpj and dados_cnpj['qsa']:
            for socio in dados_cnpj['qsa']:
                writer.writerow(["NOME", socio.get('nome', 'Não encontrado')])
                writer.writerow(["QUALIFICAÇÃO", socio.get('qual', 'Não encontrado')])
                data_entrada = socio.get('data_entrada', None)
                if data_entrada:
                    try:
                        data_entrada = datetime.strptime(data_entrada, '%d/%m/%Y').strftime('%d/%m/%Y')
                        writer.writerow(["ENTRADA", data_entrada])
                    except ValueError:
                        writer.writerow(["ENTRADA", "Data inválida"])
        else:
            writer.writerow(["", "Não encontrado"])

def salvar_json(dados_cnpj, filename):
    with open(filename, 'w', encoding='utf-8') as f:
        json.dump(dados_cnpj, f, indent=4, ensure_ascii=False)

def salvar_html(dados_cnpj, filename):
    html_content = f"""
    <!DOCTYPE html>
    <html lang="pt-BR">
    <head>
        <meta charset="UTF-8">
        <title>Consulta CNPJ</title>
        <style>
            body {{ font-family: Arial, sans-serif; margin: 20px; }}
            h1 {{ text-align: center; }}
            h2 {{ color: #333; }}
            p {{ margin: 5px 0; }}
        </style>
    </head>
    <body>
        <h1>Consulta CNPJ</h1>
        <h2>Informações Gerais</h2>
        <p><b>CNPJ:</b> {dados_cnpj.get('cnpj', 'Não encontrado')}</p>
        <p><b>RAZÃO SOCIAL:</b> {dados_cnpj.get('nome', 'Não encontrado')}</p>
        <p><b>MATRIZ OU FILIAL:</b> {dados_cnpj.get('tipo', 'Não encontrado')}</p>
        <p><b>NOME FANTASIA:</b> {dados_cnpj.get('fantasia', 'Não encontrado')}</p>
        <p><b>SITUAÇÃO CADASTRAL:</b> {dados_cnpj.get('situacao', 'Não encontrado')}</p>
        <p><b>DATA DA SITUAÇÃO CADASTRAL:</b> {dados_cnpj.get('data_situacao', 'Não encontrado')}</p>
        <p><b>MOTIVO DA SITUAÇÃO CADASTRAL:</b> {dados_cnpj.get('motivo_situacao', 'Não encontrado')}</p>
        <p><b>NATUREZA JURÍDICA:</b> {dados_cnpj.get('natureza_juridica', 'Não encontrado')}</p>
        <p><b>DATA DE ABERTURA:</b> {dados_cnpj.get('abertura', 'Não encontrado')}</p>
        <p><b>IDADE:</b> {calcular_idade(dados_cnpj['abertura']) if 'abertura' in dados_cnpj else 'Não encontrado'}</p>
        <p><b>PORTE (RFB):</b> {dados_cnpj.get('porte', 'Não encontrado')}</p>
        <p><b>CAPITAL SOCIAL:</b> R$ {dados_cnpj.get('capital_social', 'Não encontrado')}</p>
        <p><b>ATUALIZAÇÃO DESTA PÁGINA:</b> {dados_cnpj.get('ultima_atualizacao', 'Não encontrado')}</p>
        <h2>Localização</h2>
        <p><b>ENDEREÇO:</b> {dados_cnpj.get('logradouro', 'Não encontrado')} | Número: {dados_cnpj.get('numero', 'Não encontrado')}</p>
        <p><b>COMPLEMENTO:</b> {dados_cnpj.get('complemento', 'Não encontrado')}</p>
        <p><b>BAIRRO:</b> {dados_cnpj.get('bairro', 'Não encontrado')}</p>
        <p><b>CIDADE | ESTADO:</b> {dados_cnpj.get('municipio', 'Não encontrado')} | {dados_cnpj.get('uf', 'Não encontrado')}</p>
        <p><b>CEP:</b> {dados_cnpj.get('cep', 'Não encontrado')}</p>
        <p><b>TELEFONES:</b> {dados_cnpj.get('telefone', 'Não encontrado')}</p>
        <p><b>E-MAILS:</b> {dados_cnpj.get('email', 'Não encontrado')}</p>
        <h2>Atividade Econômica Principal</h2>
        <p><b>CÓDIGO:</b> {dados_cnpj['atividade_principal'][0]['code'] if 'atividade_principal' in dados_cnpj and dados_cnpj['atividade_principal'] else 'Não encontrado'}</p>
        <p><b>DESCRIÇÃO:</b> {dados_cnpj['atividade_principal'][0]['text'] if 'atividade_principal' in dados_cnpj and dados_cnpj['atividade_principal'] else 'Não encontrado'}</p>
        <h2>Atividades Econômicas Secundárias</h2>
    """
    if 'atividades_secundarias' in dados_cnpj and dados_cnpj['atividades_secundarias']:
        for atividade in dados_cnpj['atividades_secundarias']:
            html_content += f"<p><b>CÓDIGO:</b> {atividade['code']} | <b>DESCRIÇÃO:</b> {atividade['text']}</p>\n"
    else:
        html_content += "<p>Não encontrado</p>\n"

    html_content += """
        <h2>Quadro de Sócios e Administradores (QSA)</h2>
    """
    if 'qsa' in dados_cnpj and dados_cnpj['qsa']:
        for socio in dados_cnpj['qsa']:
            html_content += f"<p><b>NOME:</b> {socio.get('nome', 'Não encontrado')}</p>\n"
            html_content += f"<p><b>QUALIFICAÇÃO:</b> {socio.get('qual', 'Não encontrado')}</p>\n"
            data_entrada = socio.get('data_entrada', None)
            if data_entrada:
                try:
                    data_entrada = datetime.strptime(data_entrada, '%d/%m/%Y').strftime('%d/%m/%Y')
                    html_content += f"<p><b>ENTRADA:</b> {data_entrada}</p>\n"
                except ValueError:
                    html_content += "<p><b>ENTRADA:</b> Data inválida</p>\n"
    else:
        html_content += "<p>Não encontrado</p>\n"

    html_content += """
    </body>
    </html>
    """
    with open(filename, 'w', encoding='utf-8') as f:
        f.write(html_content)

def salvar_xml(dados_cnpj, filename):
    root = ET.Element("ConsultaCNPJ")
    
    geral = ET.SubElement(root, "InformacoesGerais")
    ET.SubElement(geral, "CNPJ").text = dados_cnpj.get('cnpj', 'Não encontrado')
    ET.SubElement(geral, "RazaoSocial").text = dados_cnpj.get('nome', 'Não encontrado')
    ET.SubElement(geral, "MatrizOuFilial").text = dados_cnpj.get('tipo', 'Não encontrado')
    ET.SubElement(geral, "NomeFantasia").text = dados_cnpj.get('fantasia', 'Não encontrado')
    ET.SubElement(geral, "SituacaoCadastral").text = dados_cnpj.get('situacao', 'Não encontrado')
    ET.SubElement(geral, "DataSituacaoCadastral").text = dados_cnpj.get('data_situacao', 'Não encontrado')
    ET.SubElement(geral, "MotivoSituacaoCadastral").text = dados_cnpj.get('motivo_situacao', 'Não encontrado')
    ET.SubElement(geral, "NaturezaJuridica").text = dados_cnpj.get('natureza_juridica', 'Não encontrado')
    ET.SubElement(geral, "DataAbertura").text = dados_cnpj.get('abertura', 'Não encontrado')
    ET.SubElement(geral, "Idade").text = calcular_idade(dados_cnpj['abertura']) if 'abertura' in dados_cnpj else 'Não encontrado'
    ET.SubElement(geral, "PorteRFB").text = dados_cnpj.get('porte', 'Não encontrado')
    ET.SubElement(geral, "CapitalSocial").text = f"R$ {dados_cnpj.get('capital_social', 'Não encontrado')}"
    ET.SubElement(geral, "UltimaAtualizacao").text = dados_cnpj.get('ultima_atualizacao', 'Não encontrado')

    localizacao = ET.SubElement(root, "Localizacao")
    ET.SubElement(localizacao, "Endereco").text = f"{dados_cnpj.get('logradouro', 'Não encontrado')} | Número: {dados_cnpj.get('numero', 'Não encontrado')}"
    ET.SubElement(localizacao, "Complemento").text = dados_cnpj.get('complemento', 'Não encontrado')
    ET.SubElement(localizacao, "Bairro").text = dados_cnpj.get('bairro', 'Não encontrado')
    ET.SubElement(localizacao, "CidadeEstado").text = f"{dados_cnpj.get('municipio', 'Não encontrado')} | {dados_cnpj.get('uf', 'Não encontrado')}"
    ET.SubElement(localizacao, "CEP").text = dados_cnpj.get('cep', 'Não encontrado')
    ET.SubElement(localizacao, "Telefones").text = dados_cnpj.get('telefone', 'Não encontrado')
    ET.SubElement(localizacao, "Emails").text = dados_cnpj.get('email', 'Não encontrado')

    atividade_principal = ET.SubElement(root, "AtividadeEconomicaPrincipal")
    ET.SubElement(atividade_principal, "Codigo").text = dados_cnpj['atividade_principal'][0]['code'] if 'atividade_principal' in dados_cnpj and dados_cnpj['atividade_principal'] else 'Não encontrado'
    ET.SubElement(atividade_principal, "Descricao").text = dados_cnpj['atividade_principal'][0]['text'] if 'atividade_principal' in dados_cnpj and dados_cnpj['atividade_principal'] else 'Não encontrado'

    atividades_secundarias = ET.SubElement(root, "AtividadesEconomicasSecundarias")
    if 'atividades_secundarias' in dados_cnpj and dados_cnpj['atividades_secundarias']:
        for atividade in dados_cnpj['atividades_secundarias']:
            atividade_elem = ET.SubElement(atividades_secundarias, "Atividade")
            ET.SubElement(atividade_elem, "Codigo").text = atividade['code']
            ET.SubElement(atividade_elem, "Descricao").text = atividade['text']
    else:
        ET.SubElement(atividades_secundarias, "Info").text = "Não encontrado"

    qsa = ET.SubElement(root, "QuadroSociosAdministradores")
    if 'qsa' in dados_cnpj and dados_cnpj['qsa']:
        for socio in dados_cnpj['qsa']:
            socio_elem = ET.SubElement(qsa, "Socio")
            ET.SubElement(socio_elem, "Nome").text = socio.get('nome', 'Não encontrado')
            ET.SubElement(socio_elem, "Qualificacao").text = socio.get('qual', 'Não encontrado')
            data_entrada = socio.get('data_entrada', None)
            if data_entrada:
                try:
                    data_entrada = datetime.strptime(data_entrada, '%d/%m/%Y').strftime('%d/%m/%Y')
                    ET.SubElement(socio_elem, "Entrada").text = data_entrada
                except ValueError:
                    ET.SubElement(socio_elem, "Entrada").text = "Data inválida"
    else:
        ET.SubElement(qsa, "Info").text = "Não encontrado"

    tree = ET.ElementTree(root)
    tree.write(filename)

def salvar_arquivo(dados_cnpj, formato_var):
    formato = formato_var.get()
    cnpj = dados_cnpj.get('cnpj', 'desconhecido').replace('/', '_').replace('.', '_').replace('-', '_')
    
    format_configs = {
        "Excel (.xlsx)": (".xlsx", [("Excel files", "*.xlsx")], f"consulta_cnpj_{cnpj}.xlsx", salvar_xlsx),
        "Word (.docx)": (".docx", [("Word files", "*.docx")], f"consulta_cnpj_{cnpj}.docx", salvar_docx),
        "Texto (.txt)": (".txt", [("Text files", "*.txt")], f"consulta_cnpj_{cnpj}.txt", salvar_txt),
        "CSV (.csv)": (".csv", [("CSV files", "*.csv")], f"consulta_cnpj_{cnpj}.csv", salvar_csv),
        "JSON (.json)": (".json", [("JSON files", "*.json")], f"consulta_cnpj_{cnpj}.json", salvar_json),
        "HTML (.html)": (".html", [("HTML files", "*.html")], f"consulta_cnpj_{cnpj}.html", salvar_html),
        "XML (.xml)": (".xml", [("XML files", "*.xml")], f"consulta_cnpj_{cnpj}.xml", salvar_xml)
    }

    if formato not in format_configs:
        messagebox.showerror("Erro", "Formato de arquivo inválido")
        return

    ext, filetypes, default_name, save_func = format_configs[formato]
    filename = filedialog.asksaveasfilename(
        title="Salvar como",
        defaultextension=ext,
        filetypes=filetypes,
        initialfile=default_name
    )
    
    if filename:
        try:
            save_func(dados_cnpj, filename)
            messagebox.showinfo("Sucesso", f"Arquivo Salvo\n\n{filename}")
        except Exception as e:
            messagebox.showerror("Erro ao salvar", f"Não foi possível salvar o arquivo: {str(e)}")

def consultar_e_mostrar(cnpj_entry, info_text, maps_button, salvar_button, formato_var):
    cnpj = limpar_cnpj(cnpj_entry.get())
    if not cnpj:
        messagebox.showerror("Erro", "Por favor, insira um CNPJ válido")
        return
    dados_cnpj = consultar_cnpj(cnpj)
    if dados_cnpj:
        info_text.config(state=tk.NORMAL)
        info_text.delete(1.0, tk.END)
        
        message = gerar_texto(dados_cnpj)
        info_text.insert(tk.END, message)
        info_text.config(state=tk.DISABLED)

        logradouro = dados_cnpj.get('logradouro', 'Não encontrado')
        numero = dados_cnpj.get('numero', 'Não encontrado')
        municipio = dados_cnpj.get('municipio', 'Não encontrado')
        uf = dados_cnpj.get('uf', 'Não encontrado')

        maps_button.config(state=tk.NORMAL, command=lambda: abrir_google_maps(logradouro, numero, municipio, uf))
        salvar_button.config(state=tk.NORMAL, command=lambda: salvar_arquivo(dados_cnpj, formato_var))


# ===================== CNPJ AVANÇADO =====================
def limpar_cnpj(cnpj):
    return ''.join(filter(str.isdigit, str(cnpj)))

def consultar_cnpj(cnpj):
    cnpj = limpar_cnpj(cnpj)
    if len(cnpj) != 14:
        return None
    try:
        r = requests.get(f"https://www.receitaws.com.br/v1/cnpj/{cnpj}", timeout=12)
        return r.json() if r.status_code == 200 else None
    except:
        return None

def calcular_idade(data_abertura):
    try:
        hoje = datetime.now()
        data = datetime.strptime(data_abertura, '%d/%m/%Y')
        diferenca = hoje - data
        anos = diferenca.days // 365
        meses = (diferenca.days % 365) // 30
        dias = (diferenca.days % 365) % 30
        return f"{anos} anos, {meses} meses e {dias} dias"
    except:
        return "Data inválida"

def abrir_google_maps(logradouro, numero, municipio, uf):
    endereco = f"{logradouro}, {numero}, {municipio}, {uf}"
    webbrowser.open(f"https://www.google.com/maps/search/?api=1&query={endereco.replace(' ', '+')}")

def gerar_texto(dados_cnpj):
    logradouro = dados_cnpj.get('logradouro', 'Não encontrado')
    numero = dados_cnpj.get('numero', 'Não encontrado')
    municipio = dados_cnpj.get('municipio', 'Não encontrado')
    uf = dados_cnpj.get('uf', 'Não encontrado')

    message = f"""
CNPJ: {dados_cnpj.get('cnpj', 'Não encontrado')}

RAZÃO SOCIAL: {dados_cnpj.get('nome', 'Não encontrado')}

MATRIZ OU FILIAL: {dados_cnpj.get('tipo', 'Não encontrado')}

NOME FANTASIA: {dados_cnpj.get('fantasia', 'Não encontrado')}

SITUAÇÃO CADASTRAL: {dados_cnpj.get('situacao', 'Não encontrado')}

DATA DA SITUAÇÃO CADASTRAL: {dados_cnpj.get('data_situacao', 'Não encontrado')}

MOTIVO DA SITUAÇÃO CADASTRAL: {dados_cnpj.get('motivo_situacao', 'Não encontrado')}

NATUREZA JURÍDICA: {dados_cnpj.get('natureza_juridica', 'Não encontrado')}

DATA DE ABERTURA: {dados_cnpj.get('abertura', 'Não encontrado')}

IDADE: {calcular_idade(dados_cnpj.get('abertura')) if 'abertura' in dados_cnpj else 'Não encontrado'}

PORTE (RFB): {dados_cnpj.get('porte', 'Não encontrado')}

CAPITAL SOCIAL: R$ {dados_cnpj.get('capital_social', 'Não encontrado')}

ATUALIZAÇÃO DESTA PÁGINA: {dados_cnpj.get('ultima_atualizacao', 'Não encontrado')}



LOCALIZAÇÃO
===========

ENDEREÇO: {logradouro} | Número: {numero}

COMPLEMENTO: {dados_cnpj.get('complemento', 'Não encontrado')}

BAIRRO: {dados_cnpj.get('bairro', 'Não encontrado')}

CIDADE | ESTADO: {municipio} | {uf}

CEP: {dados_cnpj.get('cep', 'Não encontrado')}

TELEFONES: {dados_cnpj.get('telefone', 'Não encontrado')}

E-MAILS: {dados_cnpj.get('email', 'Não encontrado')}



ATIVIDADE ECONÔMICA PRINCIPAL
==============================

CÓDIGO: {dados_cnpj.get('atividade_principal', [{}])[0].get('code', 'Não encontrado')}
DESCRIÇÃO: {dados_cnpj.get('atividade_principal', [{}])[0].get('text', 'Não encontrado')}


ATIVIDADES ECONÔMICAS SECUNDÁRIAS
=================================
"""
    if 'atividades_secundarias' in dados_cnpj and dados_cnpj['atividades_secundarias']:
        for atividade in dados_cnpj['atividades_secundarias']:
            message += f"CÓDIGO: {atividade['code']} | DESCRIÇÃO: {atividade['text']}\n"
    else:
        message += "Não encontrado\n"

    message += "\n\nQUADRO DE SÓCIOS E ADMINISTRADORES (QSA)\n==========================================\n"
    if 'qsa' in dados_cnpj and dados_cnpj['qsa']:
        for socio in dados_cnpj['qsa']:
            data_entrada = socio.get('data_entrada')
            message += f"NOME: {socio.get('nome', 'Não encontrado')}\n"
            message += f"QUALIFICAÇÃO: {socio.get('qual', 'Não encontrado')}\n"
            if data_entrada:
                try:
                    data_entrada_fmt = datetime.strptime(data_entrada, '%d/%m/%Y').strftime('%d/%m/%Y')
                    message += f"ENTRADA: {data_entrada_fmt}\n"
                except:
                    message += "ENTRADA: Data inválida\n"
            message += "-" * 50 + "\n"
    else:
        message += "Não encontrado\n"
    
    return message

# ===================== INTERFACE =====================
root = tk.Tk()
root.title("🔍 WHOIS AVANÇADO + CNPJ 🔎")
root.geometry("1200x850")
root.state('zoomed')
root.configure(bg="#0a0a0a")

# Frame Superior
top_frame = tk.Frame(root, bg="#0a0a0a")
top_frame.pack(fill="x", padx=10, pady=10)

tk.Label(top_frame, text="ALVO (Domínio / IP )", bg="#0a0a0a", fg="#00ff41", 
         font=("Consolas", 12, "bold")).pack(side=tk.LEFT, pady=5)

entry = tk.Entry(top_frame, font=("Consolas", 14), width=55, bg="#049e2b", fg="#000000")
entry.pack(side=tk.LEFT, padx=10, pady=5)

tk.Button(top_frame, text="🔍 BUSCAR WHOIS 🔎", font=("Consolas", 12, "bold"), 
          bg="#001a00", fg="#00ff41", command=lambda: buscar_tudo()).pack(side=tk.LEFT, padx=8, pady=10)

# Notebook
notebook = ttk.Notebook(root)
notebook.pack(fill="both", expand=True, padx=10, pady=5)

aba_whois = ttk.Frame(notebook)
aba_hist = ttk.Frame(notebook)
aba_cnpj = ttk.Frame(notebook)
aba_url = ttk.Frame(notebook)

notebook.add(aba_whois, text="WHOIS Atual")
notebook.add(aba_hist, text="Histórico WHOIS")
notebook.add(aba_cnpj, text="CNPJ Avançado")
notebook.add(aba_url, text="Informações URL")

# ===================== ABA CNPJ AVANÇADA =====================
cnpj_frame = tk.Frame(aba_cnpj, bg="#0a0a0a")
cnpj_frame.pack(fill="x", padx=10, pady=8)

tk.Label(cnpj_frame, text="CNPJ:", bg="#0a0a0a", fg="#00ff41", font=("Consolas", 12, "bold")).pack(side=tk.LEFT)

cnpj_entry = tk.Entry(cnpj_frame, font=("Consolas", 14), width=30, bg="#05A038", fg="#0a0a0a")
cnpj_entry.pack(side=tk.LEFT, padx=8)

btn_consultar = tk.Button(cnpj_frame, text="🔍 CONSULTAR CNPJ", font=("Consolas", 11, "bold"),
                          bg="#00ff41", fg="black", 
                          command=lambda: consultar_e_mostrar(
                              cnpj_entry, text_cnpj, maps_btn, salvar_btn, formato_var))
btn_consultar.pack(side=tk.LEFT, padx=5)

maps_btn = tk.Button(cnpj_frame, text="🗺️ GOOGLE MAPS", font=("Consolas", 11, "bold"),
                     bg="#00FFFF", fg="black", state=tk.DISABLED, width=18)
maps_btn.pack(side=tk.LEFT, padx=5)

salvar_btn = tk.Button(cnpj_frame, text="💾 SALVAR", font=("Consolas", 11, "bold"),
                       bg="#FF4500", fg="black", state=tk.DISABLED, width=15)
salvar_btn.pack(side=tk.LEFT, padx=5)

# ===================== BOTÃO ABRIR ARQUIVO (dentro da ABA CNPJ) =====================
btn_abrir = tk.Button(cnpj_frame, 
                      text="📂 ABRIR ARQUIVO SALVO CNPJ", 
                      command=abrir_arquivo,
                      font=HackerStyle.FONT_BOLD, 
                      bg="#00ff41", 
                      fg="black",
                      width=28)
btn_abrir.pack(side=tk.LEFT, padx=8)

# Menu de formatos
formato_var = tk.StringVar(value="Excel (.xlsx)")
formato_menu = tk.OptionMenu(cnpj_frame, formato_var, 
                            "Excel (.xlsx)", "Word (.docx)", "Texto (.txt)",
                            "CSV (.csv)", "JSON (.json)", "HTML (.html)", "XML (.xml)")
formato_menu.config(bg="#1a1a1a", fg="#00ff41")
formato_menu.pack(side=tk.RIGHT, padx=5)

# Text Areas
text_whois = ScrolledText(aba_whois, font=("Consolas", 11), bg="#000000", fg="#00ff41", wrap=tk.WORD)
text_whois.pack(fill="both", expand=True, padx=5, pady=5)

text_hist = ScrolledText(aba_hist, font=("Consolas", 11), bg="#000000", fg="#00ff41", wrap=tk.WORD)
text_hist.pack(fill="both", expand=True, padx=5, pady=5)

text_cnpj = ScrolledText(aba_cnpj, font=("Consolas", 11), bg="#000000", fg="#00ff41", wrap=tk.WORD)
text_cnpj.pack(fill="both", expand=True, padx=5, pady=5)

text_url = ScrolledText(aba_url, font=("Consolas", 11), bg="#000000", fg="#00ff41", wrap=tk.WORD)
text_url.pack(fill="both", expand=True, padx=5, pady=5)

# ===================== TAGS DE CORES =====================

text_whois.tag_configure("header", foreground="#00FF80", font=("Consolas", 13, "bold"))
text_whois.tag_configure("cnpj", foreground="#F7F5F4", font=("Consolas", 13, "bold"))
text_whois.tag_configure("email", foreground="#f59f16", font=("Consolas", 13, "bold"))
text_whois.tag_configure("dominio", foreground="#00BFFF", font=("Consolas", 13, "bold"))
text_whois.tag_configure("entidade", foreground="#F765F7", font=("Consolas", 13, "bold"))
text_whois.tag_configure("dns", foreground="#002EFC", font=("Consolas", 13, "bold"))
text_whois.tag_configure("data", foreground="#00FFFF", font=("Consolas", 13, "bold"))
text_whois.tag_configure("status", foreground="#F765F7", font=("Consolas", 13, "bold"))
text_whois.tag_configure("endereco", foreground="#f5ff2e", font=("Consolas", 13, "bold"))
text_whois.tag_configure("pessoa", foreground="#ff2e6d", font=("Consolas", 13, "bold"))

text_cnpj.tag_configure("header", foreground="#00FF80", font=("Consolas", 13, "bold"))
text_cnpj.tag_configure("cnpj", foreground="#FCFBFB", font=("Consolas", 13, "bold"))
text_cnpj.tag_configure("email", foreground="#f59f16", font=("Consolas", 13, "bold"))
text_cnpj.tag_configure("dominio", foreground="#00BFFF", font=("Consolas", 13, "bold"))
text_cnpj.tag_configure("entidade", foreground="#F765F7", font=("Consolas", 13, "bold"))
text_cnpj.tag_configure("dns", foreground="#002EFC", font=("Consolas", 13, "bold"))
text_cnpj.tag_configure("data", foreground="#00FFFF", font=("Consolas", 13, "bold"))
text_cnpj.tag_configure("status", foreground="#F765F7", font=("Consolas", 13, "bold"))
text_cnpj.tag_configure("endereco", foreground="#f5ff2e", font=("Consolas", 13, "bold"))
text_cnpj.tag_configure("pessoa", foreground="#ff2e6d", font=("Consolas", 13, "bold"))

footer = tk.Label(root, text="WHOIS • Consulta Segura • Informações de Registro Público + CNPJ Avançado",
                  font=("Consolas", 9), fg="#008800", bg="#0a0a0a")
footer.pack(side=tk.BOTTOM, pady=8)

# ===================== FUNÇÃO CONSULTA URL =====================
def consultar_url_info(url):
    if not url.startswith("http"):
        url = "https://" + url
    try:
        r = requests.get(url, timeout=10, allow_redirects=True)
        titulo = re.search(r'<title>(.*?)</title>', r.text, re.I | re.DOTALL)
        return f"""
        
STATUS: {r.status_code}

URL FINAL: {r.url}

TÍTULO: {titulo.group(1).strip() if titulo else 'Não encontrado'}

SERVIDOR: {r.headers.get('Server', 'N/A')}

ENCODING: {r.encoding}
        """.strip()
    except Exception as e:
        return f"Erro ao acessar URL: {e}"

# ===================== FUNÇÃO BUSCAR TUDO =====================
def buscar_tudo():
    alvo = entry.get().strip()
    if not alvo:
        messagebox.showerror("Erro", "Digite um alvo!")
        return

    # ===================== WHOIS COM CORES =====================

    text_whois.delete(1.0, tk.END)
    resultado = consultar_whois(alvo)

    for linha in resultado.splitlines():

        if ":" in linha:

            campo, valor = linha.split(":", 1)
            campo_lower = campo.lower()

            text_whois.insert("end", campo + ": ")

            if any(x in campo_lower for x in ["domínio", "domain"]):
                text_whois.insert("end", valor + "\n", "dominio")

            elif "pessoa" in campo_lower:
                text_whois.insert("end", valor + "\n", "pessoa")

            elif any(x in campo_lower for x in [
                "entidade",
                "owner",
                "registrante",
                "responsável",
                "person",
                "organização"
            ]):
                text_whois.insert("end", valor + "\n", "entidade")

            elif any(x in campo_lower for x in ["cnpj", "ownerid"]):
                text_whois.insert("end", valor + "\n", "cnpj")

            elif any(x in campo_lower for x in ["email", "e-mail", "abuse-mailbox"]):
                text_whois.insert("end", valor + "\n", "email")

            elif any(x in campo_lower for x in [
                "servidor dns",
                "servidores dns",
                "nameserver",
                "nserver",
                "dnssec"
            ]):
                text_whois.insert("end", valor + "\n", "dns")

            elif any(x in campo_lower for x in [
                "criado em",
                "alterado em",
                "expira em",
                "atualizado em",
                "creation date",
                "updated date",
                "registry expiry"
            ]):
                text_whois.insert("end", valor + "\n", "data")

            elif any(x in campo_lower for x in [
                "status",
                "status do domínio",
                "domain status"
            ]):
                text_whois.insert("end", valor + "\n", "status")

            elif any(x in campo_lower for x in [
                "endereço",
                "cidade",
                "estado",
                "cep",
                "país",
                "address",
                "country"
            ]):
                text_whois.insert("end", valor + "\n", "endereco")

            else:
                text_whois.insert("end", valor + "\n")

        else:

            if (
                "WHOIS" in linha
                or linha.startswith("=")
                or linha.startswith("-")
            ):
                text_whois.insert("end", linha + "\n", "header")
            else:
                text_whois.insert("end", linha + "\n")

    # Histórico
    text_hist.delete(1.0, tk.END)
    text_hist.insert(tk.END, f"HISTÓRICO WHOIS\n\nPara ver o histórico completo acesse\n\nhttps://www.whoxy.com/{alvo}")

    # CNPJ
    text_cnpj.delete(1.0, tk.END)
    if len(limpar_cnpj(alvo)) == 14:
        dados = consultar_cnpj(alvo)
        if dados:
            text_cnpj.insert(tk.END, gerar_texto(dados))
        else:
            text_cnpj.insert(tk.END, "CNPJ não encontrado ou API indisponível.")
    else:
        text_cnpj.insert(tk.END, "Digite um CNPJ válido (14 dígitos).")

    # URL
    text_url.delete(1.0, tk.END)
    text_url.insert(tk.END, consultar_url_info(alvo))

# Duplo clique para abrir links
def abrir_link(event):
    widget = event.widget
    try:
        index = widget.index(f"@{event.x},{event.y}")
        line = widget.get(f"{index} linestart", f"{index} lineend")
        urls = re.findall(r'https?://\S+', line)
        if urls:
            webbrowser.open(urls[0])
    except:
        pass

for txt in [text_whois, text_hist, text_cnpj, text_url]:
    txt.bind("<Double-Button-1>", abrir_link)

root.mainloop()
