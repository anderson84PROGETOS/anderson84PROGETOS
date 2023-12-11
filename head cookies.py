import tkinter as tk
from tkinter import filedialog
import requests
from urllib.parse import urlparse
from fpdf import FPDF
from docx import Document
import http.cookiejar

def salvar_cookies():
    # Obtenha a URL do campo de entrada
    url = url_entry.get()

    # Criar um objeto de gerenciamento de cookies
    cookie_jar = http.cookiejar.CookieJar()

    # Criar uma sessão HTTP usando o cookie jar
    session = requests.Session()
    session.cookies = cookie_jar

    # Defina o User-Agent para simular um navegador Windows
    user_agent = "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
    headers = {'User-Agent': user_agent}

    try:
        # Faça uma solicitação HTTP ou HTTPS para o site com o User-Agent definido
        response = session.get(url, headers=headers)

        # Obtenha o nome do site da URL
        site_name = urlparse(url).hostname

        # Obtenha o cabeçalho da página (head)
        page_head = response.headers

        # Obtenha os cookies da sessão (incluindo 'Set-Cookie' no cabeçalho)
        cookies = {}
        for cookie in session.cookies:
            cookies[cookie.name] = cookie.value

        # Solicitar ao usuário que escolha onde salvar o arquivo e em qual formato
        file_path = filedialog.asksaveasfilename(filetypes=[("Arquivos de Texto", "*.txt"),
                                                            ("Arquivos PDF", "*.pdf"),
                                                            ("Arquivos Word", "*.docx")])

        if file_path:
            if file_path.endswith(".txt"):
                with open(file_path, 'w') as f:
                    f.write(f'Site: {site_name}\n\n')
                    for cookie_name, cookie_value in cookies.items():
                        f.write(f'Cookie: {cookie_name}\n')
                        f.write(f'{cookie_name}: {cookie_value}\n\n')
                    f.write('Cabeçalho da Página (head):\n\n')
                    for key, value in page_head.items():
                        f.write(f'[{key}] {value}\n')
               
            elif file_path.endswith(".pdf"):
                pdf = FPDF()
                pdf.add_page()
                pdf.set_font("Arial", size=12)
                pdf.cell(200, 10, txt=f"Site: {site_name}", ln=True)
                for cookie_name, cookie_value in cookies.items():
                    pdf.cell(200, 10, txt=f'Cookie: {cookie_name}', ln=True)
                    pdf.cell(200, 10, txt=f'{cookie_name}: {cookie_value}', ln=True)
                pdf.cell(200, 10, txt='Cabeçalho da Página (head):', ln=True)
                for key, value in page_head.items():
                    pdf.cell(200, 10, txt=f'[{key}] {value}', ln=True)
                pdf.output(file_path)

            elif file_path.endswith(".docx"):
                doc = Document()
                doc.add_heading(f'Site: {site_name}', 0)
                for cookie_name, cookie_value in cookies.items():
                    doc.add_paragraph(f'Cookie: {cookie_name}')
                    doc.add_paragraph(f'{cookie_name}: {cookie_value}')
                    doc.add_paragraph('')
                doc.add_heading('Cabeçalho da Página (head):', level=1)
                for key, value in page_head.items():
                    doc.add_paragraph(f'[{key}] {value}')
                doc.save(file_path)

            result_label.config(text=f"Cookies salvos em {file_path}")

    except requests.exceptions.RequestException:
        result_label.config(text="Erro ao acessar a URL")

# Atualize a função escaniar_cookies com as seguintes alterações
def escaniar_cookies():
    # Obtenha a URL do campo de entrada
    url = url_entry.get()

    try:
        # Faça uma solicitação HTTP ou HTTPS para o site com um User-Agent simulando um navegador
        user_agent = "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
        headers = {'User-Agent': user_agent}
        response = requests.get(url, headers=headers)
        site_name = urlparse(url).hostname

        # Obtenha o cabeçalho da página (head) e o conteúdo da página
        page_head = response.headers
        page_content = response.text

        # Atualize o campo de texto com informações do site
        cookies_text.config(state=tk.NORMAL)
        cookies_text.delete(1.0, tk.END)
        cookies_text.insert(tk.END, f'Site: {site_name}\n\n', "site_name")
        cookies_text.insert(tk.END, 'Cabeçalho da Página (head):\n', "head_label")
        for key, value in page_head.items():
            cookies_text.insert(tk.END, f'[{key}] {value}\n', "head_value")
        cookies_text.insert(tk.END, '\nConteúdo da Página (parte inicial):\n', "head_label")

        # Limitar a exibição do conteúdo ou exibir apenas uma parte inicial
        max_content_length = 980  # Ajuste conforme necessário
        truncated_content = page_content[:max_content_length]

        cookies_text.insert(tk.END, f'{truncated_content}\n', "head_value")
        cookies_text.config(state=tk.DISABLED)

    except requests.exceptions.RequestException:
        result_label.config(text="Erro ao acessar a URL")

def criar_botao_escaniar(*args):
    url = url_entry.get()
    if url.startswith("http://") or url.startswith("https://"):
        escaniar_button.config(state=tk.NORMAL)
        save_button.config(state=tk.NORMAL)
    else:
        escaniar_button.config(state=tk.DISABLED)
        save_button.config(state=tk.DISABLED)

# Configuração da janela
window = tk.Tk()
window.wm_state('zoomed')
window.title("head_cookies")

# Rótulo de instrução
instruction_label = tk.Label(window, text="Digite a URL do site:")
instruction_label.pack()

# Campo de entrada da URL
url_entry = tk.Entry(window, width=40, fg='black', font=('TkDefaultFont', 11, 'bold'))
url_entry.pack()

# Verificar e criar botão apenas quando a URL estiver correta
url_entry.bind("<KeyRelease>", criar_botao_escaniar)

# Botão para escanear URL
escaniar_button = tk.Button(window, text="Escanear head_cookies", command=escaniar_cookies, state=tk.DISABLED, bg="#0cf2e3")
escaniar_button.pack(pady=10)

# Botão para salvar cookies
save_button = tk.Button(window, text="Salvar Cookies", command=salvar_cookies, state=tk.DISABLED, bg="#27f20c")
save_button.pack(pady=10)

# Rótulo de resultado
result_label = tk.Label(window, text="")
result_label.pack(pady=10)

# Campo de texto para mostrar cookies salvos
cookies_text = tk.Text(window, height=45, width=150, state=tk.DISABLED)
cookies_text.tag_configure("site_name", foreground="red")
cookies_text.tag_configure("cookie_label", foreground="red")
cookies_text.tag_configure("cookie_value", foreground="black")
cookies_text.tag_configure("newline", spacing1=6)
cookies_text.tag_configure("head_label", foreground="blue", font=("Arial", 12, "bold"))
cookies_text.tag_configure("head_value", foreground="black")
cookies_text.pack()

window.mainloop()
