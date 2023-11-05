import tkinter as tk
from tkinter import filedialog
import requests
import http.cookiejar
from urllib.parse import urlparse
from fpdf import FPDF
from docx import Document

def salvar_cookies():
    # Obtenha a URL do campo de entrada
    url = url_entry.get()

    # Criar um objeto de gerenciamento de cookies
    cookie_jar = http.cookiejar.CookieJar()

    # Criar uma sessão HTTP usando o cookie jar
    session = requests.Session()
    session.cookies = cookie_jar

    # Defina o User-Agent para simular um navegador Linux
    user_agent = "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/58.0.3029.110 Safari/537.36"
    headers = {'User-Agent': user_agent}

    try:
        # Faça uma solicitação HTTP ou HTTPS para o site com o User-Agent definido
        response = session.get(url, headers=headers)

        # Obtenha o nome do site da URL
        site_name = urlparse(url).hostname

        # Obtenha o cabeçalho da página (head)
        page_head = response.headers

        # Solicitar ao usuário que escolha onde salvar o arquivo e em qual formato
        file_path = filedialog.asksaveasfilename(filetypes=[("Arquivos de Texto", "*.txt"),
                                                            ("Arquivos PDF", "*.pdf"),
                                                            ("Arquivos Word", "*.docx")])

        if file_path:
            if file_path.endswith(".txt"):
                with open(file_path, 'w') as f:
                    f.write(f'Site: {site_name}\n\n')
                    for cookie in cookie_jar:
                        f.write(f'Cookie: {cookie.name}\n\n')
                        f.write(f'{cookie.name}: {cookie.value}\n\n')
                    f.write('Cabeçalho da Página (head):\n\n')
                    for key, value in page_head.items():
                        f.write(f'[{key}] {value}\n')
            elif file_path.endswith(".pdf"):
                pdf = FPDF()
                pdf.add_page()
                pdf.set_font("Arial", size=12)
                pdf.cell(200, 10, txt=f"Site: {site_name}", ln=True)
                for cookie in cookie_jar:
                    pdf.cell(200, 10, txt=f"Cookie: {cookie.name}", ln=True)
                    pdf.cell(200, 10, txt=f"{cookie.name}: {cookie.value}", ln=True)
                pdf.cell(200, 10, txt='Cabeçalho da Página (head):', ln=True)
                for key, value in page_head.items():
                    pdf.cell(200, 10, txt=f'[{key}] {value}', ln=True)
                pdf.output(file_path)
            elif file_path.endswith(".docx"):
                doc = Document()
                doc.add_heading(f'Site: {site_name}', 0)
                for cookie in cookie_jar:
                    doc.add_paragraph(f'Cookie: {cookie.name}')
                    doc.add_paragraph(f'{cookie.name}: {cookie.value}')
                doc.add_heading('Cabeçalho da Página (head):', level=1)
                for key, value in page_head.items():
                    doc.add_paragraph(f'[{key}] {value}')
                doc.save(file_path)

            result_label.config(text=f"Cookies salvos em {file_path}")
            
            # Atualize o campo de texto com os cookies salvos
            cookies_text.config(state=tk.NORMAL)
            cookies_text.delete(1.0, tk.END)
            cookies_text.insert(tk.END, f'Site: {site_name}\n', "site_name")
            for cookie in cookie_jar:
                cookies_text.insert(tk.END, '\n\n', "newline")
                cookies_text.insert(tk.END, f'Cookie: {cookie.name}\n\n', "cookie_label")
                cookies_text.insert(tk.END, f'{cookie.name}: {cookie.value}\n\n', "cookie_value")
            cookies_text.insert(tk.END, 'Cabeçalho da Página (head):\n', "head_label")
            for key, value in page_head.items():
                cookies_text.insert(tk.END, f'[{key}] {value}\n', "head_value")
            cookies_text.config(state=tk.DISABLED)

    except requests.exceptions.RequestException:
        result_label.config(text="Erro ao acessar a URL")

def criar_botao_salvar():
    url = url_entry.get()
    if url.startswith("http://") or url.startswith("https://"):
        save_button.config(state=tk.NORMAL)
    else:
        save_button.config(state=tk.DISABLED)

# Configuração da janela
window = tk.Tk()
window.wm_state('zoomed')
window.title("head_cookies")

# Rótulo de instrução
instruction_label = tk.Label(window, text="Digite a URL do site:")
instruction_label.pack()

# Campo de entrada da URL
url_entry = tk.Entry(window, width=40, fg='black', font=('TkDefaultFont', 11, 'bold'))
url_entry.pack()

# Verificar e criar botão apenas quando a URL estiver correta
url_entry.bind("<KeyRelease>", lambda event: criar_botao_salvar())

# Botão para salvar cookies
save_button = tk.Button(window, text="Salvar Cookies", command=salvar_cookies, state=tk.DISABLED)
save_button.pack()

# Rótulo de resultado
result_label = tk.Label(window, text="")
result_label.pack()

# Campo de texto para mostrar cookies salvos
cookies_text = tk.Text(window, height=50, width=150, state=tk.DISABLED)
cookies_text.tag_configure("site_name", foreground="red")
cookies_text.tag_configure("cookie_label", foreground="red")
cookies_text.tag_configure("cookie_value", foreground="black")
cookies_text.tag_configure("newline", spacing1=6)
cookies_text.tag_configure("head_label", foreground="blue", font=("Arial", 12, "bold"))
cookies_text.tag_configure("head_value", foreground="black")
cookies_text.pack()
window.mainloop()
